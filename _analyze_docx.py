"""
Разовый анализатор структуры docx: стили, разрывы страниц, шрифт, нумерация.
Удаляется после использования.
"""
import sys
import re
from docx import Document
from docx.oxml.ns import qn


def describe_run_font(run):
    name = run.font.name or "-"
    size = run.font.size.pt if run.font.size else None
    bold = bool(run.bold)
    italic = bool(run.italic)
    upper = run.text.isupper() and any(c.isalpha() for c in run.text)
    return f"{name}/{size}pt/{'B' if bold else ''}{'I' if italic else ''}{'UP' if upper else ''}"


def para_has_page_break_before(p):
    pPr = p._p.find(qn("w:pPr"))
    if pPr is None:
        return False
    pbb = pPr.find(qn("w:pageBreakBefore"))
    return pbb is not None


def run_has_page_break(r):
    for br in r._element.iter(qn("w:br")):
        t = br.get(qn("w:type"))
        if t == "page":
            return True
    return False


def para_has_toc_field(p):
    for instr in p._p.iter(qn("w:instrText")):
        if instr.text and "TOC" in instr.text:
            return True
    for fld in p._p.iter(qn("w:fldSimple")):
        instr = fld.get(qn("w:instr")) or ""
        if "TOC" in instr:
            return True
    return False


def section_info(sec):
    return {
        "page_width_cm": round(sec.page_width.cm, 2) if sec.page_width else None,
        "page_height_cm": round(sec.page_height.cm, 2) if sec.page_height else None,
        "margin_left_cm": round(sec.left_margin.cm, 2) if sec.left_margin else None,
        "margin_right_cm": round(sec.right_margin.cm, 2) if sec.right_margin else None,
        "margin_top_cm": round(sec.top_margin.cm, 2) if sec.top_margin else None,
        "margin_bottom_cm": round(sec.bottom_margin.cm, 2) if sec.bottom_margin else None,
    }


def analyze(path):
    print("=" * 90)
    print(f"ФАЙЛ: {path}")
    print("=" * 90)
    doc = Document(path)

    print("\n-- Секции --")
    for i, sec in enumerate(doc.sections):
        print(f"  section {i}: {section_info(sec)}")

    print("\n-- Стили используемые --")
    used_styles = {}
    for p in doc.paragraphs:
        s = p.style.name if p.style else "-"
        used_styles[s] = used_styles.get(s, 0) + 1
    for s, c in sorted(used_styles.items(), key=lambda x: -x[1]):
        print(f"  {s:40s} : {c}")

    print("\n-- Таблиц в документе --")
    print(f"  tables: {len(doc.tables)}")

    print("\n-- Абзацы (кратко) --")
    print(f"  paragraph count: {len(doc.paragraphs)}")

    print("\n-- Подробный обход (первые 200 абзацев) --")
    toc_found = False
    page_breaks = []
    for i, p in enumerate(doc.paragraphs[:250]):
        text = (p.text or "").replace("\n", "\\n").strip()
        pbb = para_has_page_break_before(p)
        run_pb = any(run_has_page_break(r) for r in p.runs)
        style = p.style.name if p.style else "-"
        align = p.paragraph_format.alignment
        align_s = {0: "L", 1: "C", 2: "R", 3: "J"}.get(
            align if align is not None else -1, "-"
        )
        li = p.paragraph_format.left_indent
        fli = p.paragraph_format.first_line_indent
        ls = p.paragraph_format.line_spacing

        fonts = set()
        for r in p.runs:
            if r.text:
                fonts.add(describe_run_font(r))

        has_toc = para_has_toc_field(p)
        if has_toc:
            toc_found = True

        markers = []
        if pbb:
            markers.append("PBB")
            page_breaks.append((i, "pPr", text[:60]))
        if run_pb:
            markers.append("RUN_PB")
            page_breaks.append((i, "run_br", text[:60]))
        if has_toc:
            markers.append("TOC_FIELD")

        mark = " ".join(markers)
        info_bits = []
        info_bits.append(f"style={style}")
        info_bits.append(f"al={align_s}")
        if li is not None:
            info_bits.append(f"li={round(li.cm,2)}cm")
        if fli is not None:
            info_bits.append(f"fli={round(fli.cm,2)}cm")
        if ls is not None:
            info_bits.append(f"ls={ls}")
        if fonts:
            info_bits.append("fonts=[" + ",".join(sorted(fonts))[:80] + "]")

        print(f"  [{i:03d}] {mark:20s} {'|'.join(info_bits)}")
        show = text[:120] + ("…" if len(text) > 120 else "")
        print(f"        TEXT: {show}")

    print("\n-- Сводка разрывов страницы --")
    if not page_breaks:
        print("  НЕТ ЯВНЫХ РАЗРЫВОВ СТРАНИЦ (ни pageBreakBefore, ни w:br type=page)")
    else:
        for i, kind, txt in page_breaks:
            print(f"  para[{i}] via {kind:8s}: {txt}")

    print("\n-- Поле TOC --")
    print(f"  found: {toc_found}")

    print("\n-- Поиск явных потенциальных заголовков разделов --")
    patterns = [
        r"^\s*аннотация\s*$",
        r"^\s*оглавление\s*$",
        r"^\s*содержание\s*$",
        r"^\s*введение\s*$",
        r"^\s*заключение\s*$",
        r"^\s*список\s+(использованных\s+)?источников",
        r"^\s*список\s+литературы",
        r"^\s*приложени[ея]",
        r"^\s*\d+\s+\S",
        r"^\s*\d+\.\d+\s+\S",
    ]
    for i, p in enumerate(doc.paragraphs):
        t = (p.text or "").strip()
        if not t:
            continue
        low = t.lower()
        for pat in patterns:
            if re.match(pat, low):
                pbb = para_has_page_break_before(p)
                run_pb = any(run_has_page_break(r) for r in p.runs)
                style = p.style.name if p.style else "-"
                fonts = {describe_run_font(r) for r in p.runs if r.text}
                mark = "PBB" if pbb else ("RUN_PB" if run_pb else "—no-break—")
                print(
                    f"  [{i:03d}] {mark:10s} style={style:15s} "
                    f"fonts={sorted(fonts)} :: {t[:80]}"
                )
                break


if __name__ == "__main__":
    out_path = None
    args = list(sys.argv[1:])
    if len(args) >= 2 and args[0] == "--out":
        out_path = args[1]
        args = args[2:]
    if out_path:
        f = open(out_path, "w", encoding="utf-8")
        old = sys.stdout
        sys.stdout = f
        try:
            for pth in args:
                try:
                    analyze(pth)
                except Exception as e:
                    import traceback
                    print(f"ERROR on {pth}: {e}")
                    print(traceback.format_exc())
                print("\n\n")
        finally:
            sys.stdout = old
            f.close()
        print(f"OK -> {out_path}")
    else:
        for pth in args:
            try:
                analyze(pth)
            except Exception as e:
                print(f"ERROR on {pth}: {e}")
            print("\n\n")
