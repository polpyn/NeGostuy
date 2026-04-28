"""
Форматтер курсовых работ по правилам учебного заведения.

Правила:
  • Шрифт: Times New Roman 14pt, межстрочный интервал 1,5
  • Поля: левое 3 см, правое 1,5 см, верх/низ 2 см
  • Отступ первой строки: 1,25 см
  • Выравнивание основного текста: по ширине

Заголовки 1-го уровня (special_heading / toc_heading / chapter_heading):
  — всегда начинаются с новой страницы (pageBreakBefore)

  special_heading  (ВВЕДЕНИЕ, ЗАКЛЮЧЕНИЕ, АННОТАЦИЯ, ПРИЛОЖЕНИЕ …):
    • по центру, ВСЕ ЗАГЛАВНЫЕ, жирный, Heading 1

  toc_heading  (СОДЕРЖАНИЕ / ОГЛАВЛЕНИЕ):
    • по центру, ВСЕ ЗАГЛАВНЫЕ, жирный + поле TOC

  chapter_heading  («1. Название раздела»):
    • по ширине, красная строка 1,25 см, жирный, Heading 1

Заголовки 2-го уровня (section_heading — «1.1 …»):
  • по ширине, красная строка 1,25 см, жирный, Heading 2

Заголовки 3-го уровня (subsection_heading — «1.1.1 …»):
  • по ширине, красная строка 1,25 см, жирный (курсив), Heading 3

Список (list_item): тире «— », по ширине, красная строка 1,25 см

Таблицы: стиль Table Grid, 12pt
Подписи к рисункам: 12pt, по центру, «Рисунок N – Название»
"""

import io
import os
import re
import shutil
import tempfile
import zipfile

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree


# ================================================================
# КОНСТАНТЫ
# ================================================================

KURSOVAYA = {
    "font":          "Times New Roman",
    "size":          Pt(14),
    "size_table":    Pt(12),
    "size_caption":  Pt(12),
    "line_spacing":  1.5,
    "indent":        Cm(1.25),
    "margin_left":   Cm(3.0),
    "margin_right":  Cm(1.5),
    "margin_top":    Cm(2.0),
    "margin_bottom": Cm(2.0),
}

# Разделы библиографии — в них list_item выводится с заглавной буквы и без «— »
BIB_SECTION_KEYS = {
    "библиографическое описание",
    "библиографический список",
    "список литературы",
    "список использованных источников",
    "список источников",
    "список использованной литературы",
}

ABBREV_SECTION_KEYS = {
    "перечень сокращений",
    "сокращения",
    "обозначения",
    "определения",
    "нормативные ссылки",
    "термины и определения",
}


# ================================================================
# ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ
# ================================================================

def _set_font(run, bold: bool = False, italic: bool = False, size=None):
    """Устанавливает шрифт TNR нужного размера для run."""
    run.font.name = KURSOVAYA["font"]
    run.font.size = size or KURSOVAYA["size"]
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)
    try:
        rPr = run._element.get_or_add_rPr()
        fonts_el = rPr.find(qn("w:rFonts"))
        if fonts_el is None:
            fonts_el = OxmlElement("w:rFonts")
            rPr.insert(0, fonts_el)
        fonts_el.set(qn("w:eastAsia"), KURSOVAYA["font"])
        fonts_el.set(qn("w:cs"), KURSOVAYA["font"])
    except Exception:
        pass


def _para_fmt(para,
              alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
              line_spacing: float = 1.5,
              first_line_indent=None,
              space_before: float = 0,
              space_after: float = 0):
    """Устанавливает форматирование абзаца."""
    fmt = para.paragraph_format
    fmt.alignment = alignment
    fmt.line_spacing = line_spacing
    fmt.first_line_indent = first_line_indent
    fmt.space_before = Pt(space_before)
    fmt.space_after = Pt(space_after)


def _insert_explicit_page_break(doc: Document) -> None:
    """
    Вставляет явный разрыв страницы: отдельный пустой абзац с <w:br w:type="page"/>.

    Надёжнее, чем pageBreakBefore на pPr: Word учитывает явный br-символ при подсчёте
    страниц во время обновления TOC. При pageBreakBefore Word иногда игнорирует разрыв
    на этапе field-update и нумерует все заголовки одной страницей.
    """
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    r = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r.append(br)
    para._p.append(r)


def _clear_heading_style_numbering(para):
    """
    Убирает автоматическую нумерацию из стиля Heading 1/2/3,
    чтобы слово «ВВЕДЕНИЕ» не превращалось в «1 ВВЕДЕНИЕ».
    """
    pPr = para._element.get_or_add_pPr()
    numPr = pPr.find(qn("w:numPr"))
    if numPr is not None:
        pPr.remove(numPr)


def _clear_paragraph_tabs_and_numbering(para) -> None:
    """
    Удаляет таб-стопы и numPr у абзаца. Это лечит «огромные отступы»,
    которые появляются из-за унаследованных tabs/нумерации в стилях Word.
    """
    pPr = para._element.get_or_add_pPr()
    tabs = pPr.find(qn("w:tabs"))
    if tabs is not None:
        pPr.remove(tabs)
    numPr = pPr.find(qn("w:numPr"))
    if numPr is not None:
        pPr.remove(numPr)


_LEADING_JUNK_RE = re.compile(r"^[\s\u00A0\t]+")
_LEADING_MARKER_RE = re.compile(r"^(?:[—\-–•●○]\s+)+")


def _clean_entry_text(text: str) -> str:
    """
    Нормализует начало строки записи (библиография/сокращения):
    - убирает лидирующие табы/неразрывные пробелы
    - убирает маркеры списка «—/•» в начале
    - схлопывает лишние пробелы
    """
    t = (text or "")
    t = _LEADING_JUNK_RE.sub("", t)
    t = _LEADING_MARKER_RE.sub("", t)
    t = t.replace("\t", " ")
    t = re.sub(r"[ \u00A0]{2,}", " ", t).strip()
    return t


def _ensure_trailing_punct(text: str, punct: str) -> str:
    """
    Нормализует знак в конце пункта списка.
    - если уже заканчивается на punct — оставляем
    - если заканчивается на ';' или '.' — заменяем на punct
    - иначе добавляем punct
    """
    t = (text or "").rstrip()
    if not t:
        return t
    if t.endswith(punct):
        return t
    if t.endswith(";") or t.endswith("."):
        return t[:-1] + punct
    return t + punct


def _insert_toc_field(doc: Document):
    """
    Вставляет поле Table of Contents (TOC) после текущей позиции.
    Поле обновляется при первом открытии документа в Word.
    Параметры: заголовки Heading 1..3, гиперссылки, без номеров стилей.

    Структура: begin + instrText + end (без separate/placeholder),
    dirty="true" — Word пересчитает номера страниц с нуля при открытии.
    Вставка placeholder'а внутрь одного абзаца приводит к тому, что Word
    вычисляет страницу самого абзаца-TOC и подставляет её во все записи.
    """
    para = doc.add_paragraph()
    _para_fmt(para,
              alignment=WD_ALIGN_PARAGRAPH.LEFT,
              line_spacing=KURSOVAYA["line_spacing"],
              first_line_indent=Cm(0))

    # begin (dirty=true — принудительное обновление при открытии)
    r_begin = para.add_run()
    fc_begin = OxmlElement("w:fldChar")
    fc_begin.set(qn("w:fldCharType"), "begin")
    fc_begin.set(qn("w:dirty"), "true")
    r_begin._r.append(fc_begin)

    # instrText
    r_instr = para.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    r_instr._r.append(instr)

    # end (без separate — Word сам построит записи при обновлении)
    r_end = para.add_run()
    fc_end = OxmlElement("w:fldChar")
    fc_end.set(qn("w:fldCharType"), "end")
    r_end._r.append(fc_end)

    return para


def _gost_image_width(image_path: str, max_w_cm: float = 16.0, max_h_cm: float = 18.0):
    """Возвращает ширину вставки с сохранением пропорций."""
    try:
        from PIL import Image
        with Image.open(image_path) as im:
            px_w, px_h = im.size
            dpi_raw = im.info.get("dpi")
            dpi_x = float(dpi_raw[0]) if dpi_raw and dpi_raw[0] else 96.0
            dpi_y = float(dpi_raw[1]) if dpi_raw and len(dpi_raw) > 1 and dpi_raw[1] else dpi_x
        w_cm = (px_w / dpi_x) * 2.54
        h_cm = (px_h / dpi_y) * 2.54
        if w_cm > 0 and h_cm > 0:
            scale = min(max_w_cm / w_cm, max_h_cm / h_cm, 1.0)
            return Cm(w_cm * scale)
    except Exception:
        pass
    return Cm(min(max_w_cm, 16.0))


# ================================================================
# ОБНОВЛЕНИЕ ПОЛЕЙ (w:updateFields) — чтобы Word пересчитал TOC
# ================================================================

def _enable_update_fields_on_open(docx_path: str) -> None:
    """
    Дописывает <w:updateFields w:val="true"/> в word/settings.xml.
    При следующем открытии Word предложит обновить все поля (в т.ч. TOC).
    """
    WD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    SETTINGS_INNER = "word/settings.xml"

    fd, tmp_path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)

    try:
        with zipfile.ZipFile(docx_path, "r") as zin, \
             zipfile.ZipFile(tmp_path, "w", compression=zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == SETTINGS_INNER:
                    tree = etree.fromstring(data)
                    # Проверяем, нет ли уже updateFields
                    tag = f"{{{WD_NS}}}updateFields"
                    existing = tree.find(tag)
                    if existing is None:
                        uf = etree.SubElement(tree, tag)
                        uf.set(f"{{{WD_NS}}}val", "true")
                        data = etree.tostring(tree, xml_declaration=True,
                                              encoding="UTF-8", standalone=True)
                zout.writestr(item, data)
        shutil.move(tmp_path, docx_path)
    except Exception as exc:
        print(f"  ⚠️  _enable_update_fields_on_open: {exc}")
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)


# ================================================================
# ПАТЧ СТИЛЕЙ (TOC / Hyperlink) — убрать синий цвет
# ================================================================

def _patch_toc_and_hyperlink_styles(docx_path: str) -> None:
    """
    Убирает «синие ссылки» у оглавления: Word применяет стили Hyperlink и TOC1..TOC3
    при обновлении поля TOC. Патчим word/styles.xml: цвет -> чёрный, underline -> none.

    Дополнительно чиним типичный шаблон рамки: встроенные стили Heading 1..3 в styles.xml
    часто задают themeColor (синий). После LLM-постобработки прямые w:color у run'ов могут
    пропасть — тогда заголовки снова «синеют» из стиля.
    """
    WD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    STYLES_INNER = "word/styles.xml"

    fd, tmp_path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)

    def _ensure_run_props(style_el: etree._Element) -> etree._Element:
        rpr = style_el.find(f"{{{WD_NS}}}rPr")
        if rpr is None:
            rpr = etree.SubElement(style_el, f"{{{WD_NS}}}rPr")
        return rpr

    def _set_color_and_underline(rpr: etree._Element) -> None:
        color = rpr.find(f"{{{WD_NS}}}color")
        if color is None:
            color = etree.SubElement(rpr, f"{{{WD_NS}}}color")
        color.set(f"{{{WD_NS}}}val", "000000")
        # Убираем theme-атрибуты, иначе Word может продолжать красить «по теме»
        for attr in list(color.attrib):
            if attr != f"{{{WD_NS}}}val":
                del color.attrib[attr]

        u = rpr.find(f"{{{WD_NS}}}u")
        if u is None:
            u = etree.SubElement(rpr, f"{{{WD_NS}}}u")
        u.set(f"{{{WD_NS}}}val", "none")

    def _style_exists(tree: etree._Element, style_id: str) -> bool:
        for st in tree.findall(f".//{{{WD_NS}}}style"):
            if (st.get(f"{{{WD_NS}}}styleId") or "") == style_id:
                return True
        return False

    def _append_child(parent: etree._Element, tag: str, **attrs) -> etree._Element:
        el = etree.SubElement(parent, f"{{{WD_NS}}}{tag}")
        for k, v in attrs.items():
            if v is None:
                continue
            el.set(f"{{{WD_NS}}}{k}", str(v))
        return el

    def _ensure_builtin_hyperlink_and_toc_styles(tree: etree._Element) -> None:
        """
        В «лёгких» шаблонах Hyperlink/TOC* могут отсутствовать как реальные <w:style>,
        оставаясь только в <w:latentStyles>. Тогда Word берёт дефолты (синие ссылки).
        Добавляем минимальные определения стилей в конец styles.xml.
        """
        # character style Hyperlink
        if not _style_exists(tree, "Hyperlink"):
            st = _append_child(tree, "style", type="character", customStyle="1", styleId="Hyperlink")
            _append_child(st, "name", val="Hyperlink")
            _append_child(st, "uiPriority", val="99")
            _append_child(st, "unhideWhenUsed")
            _set_color_and_underline(_ensure_run_props(st))

        def _default_paragraph_style_id() -> str:
            for st in tree.findall(f".//{{{WD_NS}}}style"):
                if (st.get(f"{{{WD_NS}}}type") or "") != "paragraph":
                    continue
                d = st.get(f"{{{WD_NS}}}default")
                if d in ("1", "true", "on"):
                    sid0 = st.get(f"{{{WD_NS}}}styleId") or ""
                    if sid0:
                        return sid0
            # fallback: часто Normal = "a" в шаблонах Word
            return "a"

        base_para = _default_paragraph_style_id()

        # paragraph styles TOC 1..3 (минимально достаточно для цвета/подчёркивания)
        for sid, wname in (("TOC1", "toc 1"), ("TOC2", "toc 2"), ("TOC3", "toc 3")):
            if _style_exists(tree, sid):
                continue
            st = _append_child(tree, "style", type="paragraph", customStyle="1", styleId=sid)
            _append_child(st, "name", val=wname)
            _append_child(st, "basedOn", val=base_para)
            _append_child(st, "next", val=base_para)
            _append_child(st, "uiPriority", val="39")
            _append_child(st, "unhideWhenUsed")
            _append_child(st, "semiHidden")
            _append_child(st, "qFormat")
            _set_color_and_underline(_ensure_run_props(st))

    # Hyperlink/TOC — для оглавления; 1/2/3 и a/10/20/30 — частые styleId в шаблонах Word
    # для Normal/Heading и связанных character styles.
    target_style_ids = {
        "Hyperlink",
        "FollowedHyperlink",
        "TOC1",
        "TOC2",
        "TOC3",
        "TOCHeading",
        # Heading paragraph styles (часто в локализованных/«урезанных» шаблонах)
        "1",
        "2",
        "3",
        # Normal + linked char styles
        "a",
        "10",
        "20",
        "30",
    }

    try:
        with zipfile.ZipFile(docx_path, "r") as zin, \
             zipfile.ZipFile(tmp_path, "w", compression=zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == STYLES_INNER:
                    try:
                        tree = etree.fromstring(data)
                        _ensure_builtin_hyperlink_and_toc_styles(tree)
                        # В OOXML порядок атрибутов у <w:style> не фиксирован — ищем по styleId,
                        # а не по шаблону «styleId в начале тега».
                        for st in tree.findall(f".//{{{WD_NS}}}style"):
                            style_id = st.get(f"{{{WD_NS}}}styleId") or ""
                            if style_id in target_style_ids:
                                rpr = _ensure_run_props(st)
                                _set_color_and_underline(rpr)
                        data = etree.tostring(
                            tree,
                            xml_declaration=True,
                            encoding="UTF-8",
                            standalone=True,
                        )
                    except Exception:
                        # styles.xml может отсутствовать/быть нестандартным — пропускаем
                        pass
                zout.writestr(item, data)
        shutil.move(tmp_path, docx_path)
    except Exception as exc:
        print(f"  ⚠️  _patch_toc_and_hyperlink_styles: {exc}")
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)


# ================================================================
# ГЛАВНАЯ ФУНКЦИЯ
# ================================================================

def create_kursovaya_document(
    elements: list,
    output_path: str,
    template_path: str | None = None,
    zachet_number: str | None = None,
) -> str:
    """
    Создаёт новый .docx с правильным оформлением курсовой работы.

    Args:
        elements:       список dict из kursovaya_classifier
        output_path:    путь для сохранения результата
        template_path:  путь к рамке .docx (необязательно)
        zachet_number:  номер зачётной книжки для подстановки в штамп {{nomer}}

    Returns:
        output_path
    """
    print(f"\n{'='*60}")
    print("📚 ФОРМАТИРОВАНИЕ КУРСОВОЙ РАБОТЫ")
    print(f"{'='*60}")

    use_template = False

    if template_path and os.path.exists(template_path):
        print(f"  🖼️ Рамка: {os.path.basename(template_path)}")
        try:
            from .formatter import prepare_template
            doc = prepare_template(template_path)
            use_template = True
            print("  ✅ Рамка применена!")
        except Exception as exc:
            print(f"  ❌ Ошибка рамки: {exc}")
            doc = Document()
    else:
        if template_path:
            print(f"  ❌ Файл рамки не найден: {template_path}")
        doc = Document()

    # ── Поля страницы (только если рамка не задаёт свои) ───────────
    if not use_template:
        for sec in doc.sections:
            sec.left_margin   = KURSOVAYA["margin_left"]
            sec.right_margin  = KURSOVAYA["margin_right"]
            sec.top_margin    = KURSOVAYA["margin_top"]
            sec.bottom_margin = KURSOVAYA["margin_bottom"]

    # Счётчик — первый заголовок 1-го уровня не нуждается в разрыве
    # (иначе документ начнётся с пустой страницы)
    _first_lvl1 = True
    _written = 0
    in_bibliography = False  # True когда текущий раздел — библиографический
    in_abbrev = False        # True когда текущий раздел — перечень сокращений/обозначений
    bib_counter = 0          # Нумерация записей внутри раздела библиографии

    for elem in elements:
        etype = elem.get("type", "paragraph")
        text  = (elem.get("text") or "").strip()

        # ── TABLE ─────────────────────────────────────────────────
        if etype == "table":
            rows_data = elem.get("rows") or []
            if not rows_data:
                continue
            ncol  = max(len(r) for r in rows_data)
            nrows = len(rows_data)
            tbl = doc.add_table(rows=nrows, cols=ncol)
            try:
                tbl.style = "Table Grid"
            except (KeyError, ValueError):
                pass
            for ri, row_cells in enumerate(rows_data):
                for ci in range(ncol):
                    txt = row_cells[ci] if ci < len(row_cells) else ""
                    cell = tbl.cell(ri, ci)
                    cell.text = txt
                    for p in cell.paragraphs:
                        _para_fmt(p,
                                  alignment=WD_ALIGN_PARAGRAPH.LEFT,
                                  line_spacing=KURSOVAYA["line_spacing"],
                                  first_line_indent=Cm(0))
                        for r in p.runs:
                            _set_font(r, size=KURSOVAYA["size_table"])
            _written += 1
            continue

        # ── IMAGE ─────────────────────────────────────────────────
        if etype == "image":
            for img_path in elem.get("image_paths", []):
                if not os.path.exists(img_path):
                    continue
                para = doc.add_paragraph()
                _para_fmt(para,
                          alignment=WD_ALIGN_PARAGRAPH.CENTER,
                          line_spacing=KURSOVAYA["line_spacing"],
                          first_line_indent=Cm(0),
                          space_before=6, space_after=6)
                run = para.add_run()
                try:
                    w = _gost_image_width(img_path)
                    run.add_picture(img_path, width=w)
                    _written += 1
                except Exception as exc:
                    print(f"  ⚠️  Изображение: {exc}")
            continue

        # ── FIGURE CAPTION ────────────────────────────────────────
        if etype == "figure_caption":
            clean = re.sub(
                r"^рисунок\s+\d+\s*[–\-—.:]\s*", "", text, flags=re.IGNORECASE
            )
            fig_num = elem.get("figure_num", 1)
            para = doc.add_paragraph()
            _para_fmt(para,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER,
                      line_spacing=1.0,
                      first_line_indent=Cm(0),
                      space_before=6, space_after=12)
            run = para.add_run(f"Рисунок {fig_num} – {clean}")
            _set_font(run, size=KURSOVAYA["size_caption"])
            _written += 1
            continue

        # Все остальные типы требуют ненулевого текста
        if not text:
            continue

        # ── TOC HEADING (СОДЕРЖАНИЕ / ОГЛАВЛЕНИЕ) ────────────────
        if etype == "toc_heading":
            tl = text.strip().lower()
            in_bibliography = tl in BIB_SECTION_KEYS
            in_abbrev = tl in ABBREV_SECTION_KEYS
            bib_counter = 0
            if not _first_lvl1:
                _insert_explicit_page_break(doc)
            _first_lvl1 = False
            para = doc.add_paragraph(style="Heading 1")
            _clear_heading_style_numbering(para)
            _para_fmt(para,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER,
                      line_spacing=KURSOVAYA["line_spacing"],
                      first_line_indent=Cm(0),
                      space_before=0, space_after=12)
            run = para.add_run(text.upper())
            _set_font(run, bold=True)
            # Вставляем поле TOC
            _insert_toc_field(doc)
            _written += 1
            continue

        # ── SPECIAL HEADING (ВВЕДЕНИЕ, ЗАКЛЮЧЕНИЕ, ПРИЛОЖЕНИЕ …) ─
        if etype == "special_heading":
            tl = text.strip().lower()
            in_bibliography = tl in BIB_SECTION_KEYS
            in_abbrev = tl in ABBREV_SECTION_KEYS
            bib_counter = 0
            if not _first_lvl1:
                _insert_explicit_page_break(doc)
            _first_lvl1 = False
            para = doc.add_paragraph(style="Heading 1")
            _clear_heading_style_numbering(para)
            _para_fmt(para,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER,
                      line_spacing=KURSOVAYA["line_spacing"],
                      first_line_indent=Cm(0),
                      space_before=0, space_after=12)
            run = para.add_run(text.upper())
            _set_font(run, bold=True)
            _written += 1
            continue

        # ── CHAPTER HEADING («1. Название раздела») ──────────────
        if etype == "chapter_heading":
            in_bibliography = False
            in_abbrev = False
            bib_counter = 0
            if not _first_lvl1:
                _insert_explicit_page_break(doc)
            _first_lvl1 = False
            para = doc.add_paragraph(style="Heading 1")
            _clear_heading_style_numbering(para)
            _para_fmt(para,
                      alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                      line_spacing=KURSOVAYA["line_spacing"],
                      first_line_indent=KURSOVAYA["indent"],
                      space_before=0, space_after=12)
            run = para.add_run(text)
            _set_font(run, bold=True)
            _written += 1
            continue

        # ── SECTION HEADING («1.1 Название») ─────────────────────
        if etype == "section_heading":
            para = doc.add_paragraph(style="Heading 2")
            _clear_heading_style_numbering(para)
            _para_fmt(para,
                      alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                      line_spacing=KURSOVAYA["line_spacing"],
                      first_line_indent=KURSOVAYA["indent"],
                      space_before=12, space_after=6)
            run = para.add_run(text)
            _set_font(run, bold=True)
            _written += 1
            continue

        # ── SUBSECTION HEADING («1.1.1 Название») ────────────────
        if etype == "subsection_heading":
            para = doc.add_paragraph(style="Heading 3")
            _clear_heading_style_numbering(para)
            _para_fmt(para,
                      alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                      line_spacing=KURSOVAYA["line_spacing"],
                      first_line_indent=KURSOVAYA["indent"],
                      space_before=6, space_after=6)
            run = para.add_run(text)
            _set_font(run, bold=True)
            _written += 1
            continue

        # ── LIST ITEM ─────────────────────────────────────────────
        if etype == "list_item":
            if in_abbrev:
                # Перечень сокращений: без маркера, с красной строкой 1.25 см.
                clean = _clean_entry_text(text)
                para = doc.add_paragraph()
                _para_fmt(
                    para,
                    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                    line_spacing=KURSOVAYA["line_spacing"],
                    first_line_indent=KURSOVAYA["indent"],
                    space_before=0,
                    space_after=0,
                )
                # На всякий случай переопределяем левый отступ: если стиль/Word «подцепил»
                # список или таб-стопы, может получиться визуальный отступ ~2.5 см.
                para.paragraph_format.left_indent = Cm(0)
                _clear_paragraph_tabs_and_numbering(para)
                run = para.add_run(clean)
                _set_font(run)
                _written += 1
                continue

            if in_bibliography:
                # Библиографические записи: нумерация 1., 2., 3. + красная строка
                clean = _clean_entry_text(text)
                if not clean:
                    continue
                bib_counter += 1
                para = doc.add_paragraph()
                _para_fmt(para,
                          alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                          line_spacing=KURSOVAYA["line_spacing"],
                          first_line_indent=KURSOVAYA["indent"],
                          space_before=0, space_after=0)
                _clear_paragraph_tabs_and_numbering(para)
                run = para.add_run(f"{bib_counter}. {clean}")
                _set_font(run)
                _written += 1
                continue

            clean = re.sub(r"^[—\-–•●○]\s*", "", text).strip()
            # Сохраняем нумерованные пункты как есть, маркированные — с тире
            if re.match(r"^\d+\.", text):
                # Нумерованный пункт: в конце должна быть точка
                display = _ensure_trailing_punct(text, ".")
            else:
                # Кириллическую заглавную первую букву делаем строчной;
                # латинские имена (Telegram, Python) остаются без изменений.
                if clean and re.match(r"^[А-ЯЁ]", clean):
                    clean = clean[0].lower() + clean[1:]
                # Маркированный пункт: в конце должна быть ';'
                display = _ensure_trailing_punct("— " + clean, ";")
            para = doc.add_paragraph()
            # LEFT вместо JUSTIFY: при JUSTIFY Word растягивает пробелы на
            # первой строке двустрочного пункта — зазор между «—» и словом
            # становится непредсказуемым.
            # left_indent=1.25cm — весь абзац сдвинут на 1.25 (и тире, и перенос строки).
            # first_line_indent=0 — дополнительный отступ первой строки не нужен.
            _para_fmt(para,
                      alignment=WD_ALIGN_PARAGRAPH.LEFT,
                      line_spacing=KURSOVAYA["line_spacing"],
                      first_line_indent=Cm(0),
                      space_before=0, space_after=0)
            para.paragraph_format.left_indent = KURSOVAYA["indent"]
            _clear_paragraph_tabs_and_numbering(para)
            run = para.add_run(display)
            _set_font(run)
            _written += 1
            continue

        # ── PARAGRAPH (основной текст) ────────────────────────────
        # Внутри библиографии «paragraph» тоже считаем записью (часто без маркеров)
        if in_bibliography:
            clean = _clean_entry_text(text)
            if clean:
                bib_counter += 1
                para = doc.add_paragraph()
                _para_fmt(
                    para,
                    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                    line_spacing=KURSOVAYA["line_spacing"],
                    first_line_indent=KURSOVAYA["indent"],
                    space_before=0,
                    space_after=0,
                )
                _clear_paragraph_tabs_and_numbering(para)
                run = para.add_run(f"{bib_counter}. {clean}")
                _set_font(run)
                _written += 1
                continue

        para = doc.add_paragraph()
        _para_fmt(para,
                  alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  line_spacing=KURSOVAYA["line_spacing"],
                  first_line_indent=KURSOVAYA["indent"],
                  space_before=0, space_after=0)
        _clear_paragraph_tabs_and_numbering(para)
        run = para.add_run(_clean_entry_text(text) if in_abbrev else text)
        _set_font(run)
        _written += 1

    # ── Сохранение ────────────────────────────────────────────────
    out_dir = os.path.dirname(output_path)
    if out_dir:
        os.makedirs(out_dir, exist_ok=True)
    doc.save(output_path)

    # Подставляем {{nomer}} и {{list}} в штамп рамки
    if use_template:
        from .frame_placeholders import patch_docx_headers
        patch_docx_headers(output_path, (zachet_number or "").strip())

    # Добавляем w:updateFields, чтобы Word предложил обновить TOC при открытии
    _enable_update_fields_on_open(output_path)

    # Убираем «синий цвет ссылок» в оглавлении (Hyperlink/TOC стили Word)
    _patch_toc_and_hyperlink_styles(output_path)

    print(f"\n  ✅ Записано элементов: {_written}")
    print(f"  📂 Файл: {output_path} ({os.path.getsize(output_path):,} байт)")

    return output_path
