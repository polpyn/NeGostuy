"""
Диагностика рамки — запускай отдельно:
  python debug_template.py путь_к_рамке.docx
"""

import sys
import os
import zipfile
import shutil
from lxml import etree

# Чтобы Django не мешал
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))


def dump_docx_structure(docx_path, label=""):
    """Показывает полную структуру .docx"""
    print(f"\n{'='*70}")
    print(f"📦 СТРУКТУРА: {label or docx_path}")
    print(f"{'='*70}")

    with zipfile.ZipFile(docx_path, 'r') as z:
        print(f"\n📂 Все файлы внутри .docx:")
        for fi in sorted(z.filelist, key=lambda x: x.filename):
            print(f"   {fi.filename:50s}  ({fi.file_size} bytes)")

        # Проверяем наличие headers/footers
        headers = [f for f in z.namelist() if 'header' in f.lower()]
        footers = [f for f in z.namelist() if 'footer' in f.lower()]
        print(f"\n📋 Headers: {headers}")
        print(f"📋 Footers: {footers}")

        # Показываем содержимое [Content_Types].xml
        print(f"\n📄 [Content_Types].xml:")
        try:
            ct = z.read('[Content_Types].xml').decode('utf-8')
            for line in ct.split('<'):
                if 'header' in line.lower() or 'footer' in line.lower():
                    print(f"   <{line.strip()}")
        except:
            print("   (не удалось прочитать)")

        # Показываем word/_rels/document.xml.rels
        print(f"\n📄 word/_rels/document.xml.rels:")
        try:
            rels = z.read('word/_rels/document.xml.rels').decode('utf-8')
            for line in rels.split('<'):
                stripped = line.strip()
                if stripped:
                    print(f"   <{stripped}")
        except:
            print("   (не удалось прочитать)")

        # Показываем document.xml (sectPr часть)
        print(f"\n📄 word/document.xml (структура body):")
        try:
            doc_xml = z.read('word/document.xml').decode('utf-8')
            root = etree.fromstring(doc_xml.encode('utf-8'))
            ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

            body = root.find('.//w:body', ns)
            if body is not None:
                print(f"   Элементов в body: {len(body)}")
                for i, child in enumerate(body):
                    tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
                    if tag == 'p':
                        # Первые 50 символов текста
                        texts = child.itertext()
                        text = ''.join(texts)[:50]
                        print(f"   [{i}] <w:{tag}> text='{text}'")
                    elif tag == 'sectPr':
                        print(f"   [{i}] <w:{tag}> ⬅️ СЕКЦИЯ (рамка тут)")
                        # Показываем содержимое sectPr
                        sect_xml = etree.tostring(child, pretty_print=True).decode()
                        for line in sect_xml.split('\n')[:20]:
                            print(f"        {line}")
                    else:
                        print(f"   [{i}] <w:{tag}>")
            else:
                print("   ❌ body не найден!")
        except Exception as e:
            print(f"   Ошибка: {e}")

        # Показываем содержимое header*.xml
        for h in headers:
            if h.endswith('.xml'):
                print(f"\n📄 {h}:")
                try:
                    hxml = z.read(h).decode('utf-8')
                    root = etree.fromstring(hxml.encode('utf-8'))
                    # Проверяем наличие изображений
                    has_blip = 'blip' in hxml
                    has_pict = 'pict' in hxml
                    has_vml = 'v:' in hxml or 'vml' in hxml.lower()
                    has_shape = 'shape' in hxml.lower()
                    print(f"   blip(изображение): {has_blip}")
                    print(f"   pict: {has_pict}")
                    print(f"   VML: {has_vml}")
                    print(f"   shape: {has_shape}")
                    print(f"   Размер: {len(hxml)} символов")
                    # Первые 500 символов
                    print(f"   Начало: {hxml[:500]}")
                except Exception as e:
                    print(f"   Ошибка: {e}")

        for f in footers:
            if f.endswith('.xml'):
                print(f"\n📄 {f}:")
                try:
                    fxml = z.read(f).decode('utf-8')
                    has_blip = 'blip' in fxml
                    has_pict = 'pict' in fxml
                    has_vml = 'v:' in fxml or 'vml' in fxml.lower()
                    print(f"   blip(изображение): {has_blip}")
                    print(f"   pict: {has_pict}")
                    print(f"   VML: {has_vml}")
                    print(f"   Размер: {len(fxml)} символов")
                except Exception as e:
                    print(f"   Ошибка: {e}")


def test_template_processing(template_path):
    """Тестирует обработку рамки от начала до конца"""
    from docx import Document as DocxDocument
    from docx.oxml.ns import qn

    print(f"\n{'='*70}")
    print(f"🧪 ТЕСТ ОБРАБОТКИ РАМКИ")
    print(f"{'='*70}")

    # ШАГ 1: Открываем рамку
    print(f"\n--- ШАГ 1: Открытие рамки ---")
    doc = DocxDocument(template_path)
    body = doc.element.body

    print(f"Элементов в body: {len(body)}")
    for i, child in enumerate(body):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        print(f"  [{i}] {tag}")

    # Ищем sectPr
    sectPrs = [c for c in body if c.tag == qn('w:sectPr') or c.tag.endswith('}sectPr')]
    print(f"sectPr найдено: {len(sectPrs)}")

    # Проверяем rels документа
    print(f"\nDocument rels:")
    for rel_id, rel in doc.part.rels.items():
        if 'header' in str(rel.reltype).lower() or 'footer' in str(rel.reltype).lower():
            print(f"  {rel_id}: {rel.reltype} -> {rel.target_ref}")

    # Проверяем секции
    print(f"\nСекций: {len(doc.sections)}")
    for si, section in enumerate(doc.sections):
        print(f"  Секция {si}:")
        try:
            sect_xml = etree.tostring(section._sectPr, pretty_print=True).decode()
            for line in sect_xml.split('\n')[:15]:
                if line.strip():
                    print(f"    {line}")
        except Exception as e:
            print(f"    Ошибка: {e}")

    # ШАГ 2: Удаляем параграфы
    print(f"\n--- ШАГ 2: Удаление параграфов ---")
    removed = 0
    for child in list(body):
        if child.tag in (qn('w:p'), qn('w:tbl')):
            body.remove(child)
            removed += 1
    print(f"Удалено: {removed}")
    print(f"Осталось в body: {len(body)}")
    for i, child in enumerate(body):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        print(f"  [{i}] {tag}")

    # ШАГ 3: Добавляем тестовый параграф
    print(f"\n--- ШАГ 3: Добавление тестового параграфа ---")
    para = doc.add_paragraph("Тестовый параграф для проверки рамки")
    print(f"Элементов в body после add_paragraph: {len(body)}")
    for i, child in enumerate(body):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'p':
            texts = list(child.itertext())
            text = ''.join(texts)[:50]
            print(f"  [{i}] {tag}: '{text}'")
        else:
            print(f"  [{i}] {tag}")

    # ШАГ 4: Перемещаем sectPr в конец
    print(f"\n--- ШАГ 4: Перемещение sectPr ---")
    for child in list(body):
        if child.tag == qn('w:sectPr') or child.tag.endswith('}sectPr'):
            body.remove(child)
            body.append(child)
            print(f"  sectPr перемещён в конец")

    print(f"Финальная структура body ({len(body)} элементов):")
    for i, child in enumerate(body):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'p':
            texts = list(child.itertext())
            text = ''.join(texts)[:50]
            print(f"  [{i}] {tag}: '{text}'")
        else:
            print(f"  [{i}] {tag}")

    # ШАГ 5: Сохраняем
    print(f"\n--- ШАГ 5: Сохранение ---")
    output = template_path.replace('.docx', '_TEST_OUTPUT.docx')
    doc.save(output)
    print(f"Сохранено: {output}")
    print(f"Размер: {os.path.getsize(output)} байт")

    # ШАГ 6: Проверяем результат
    print(f"\n--- ШАГ 6: Проверка результата ---")
    dump_docx_structure(output, "РЕЗУЛЬТАТ")

    # Сравниваем rels
    print(f"\n--- СРАВНЕНИЕ RELS ---")
    print(f"Оригинал:")
    with zipfile.ZipFile(template_path, 'r') as z:
        try:
            print(z.read('word/_rels/document.xml.rels').decode('utf-8')[:1000])
        except:
            pass

    print(f"\nРезультат:")
    with zipfile.ZipFile(output, 'r') as z:
        try:
            print(z.read('word/_rels/document.xml.rels').decode('utf-8')[:1000])
        except:
            pass

    # Проверяем что headers скопировались
    print(f"\n--- НАЛИЧИЕ ФАЙЛОВ ---")
    with zipfile.ZipFile(template_path, 'r') as z1:
        with zipfile.ZipFile(output, 'r') as z2:
            orig_files = set(z1.namelist())
            out_files = set(z2.namelist())

            missing = orig_files - out_files
            extra = out_files - orig_files

            if missing:
                print(f"❌ ОТСУТСТВУЮТ в результате:")
                for f in sorted(missing):
                    print(f"   {f}")
            else:
                print(f"✅ Все файлы из оригинала есть в результате")

            if extra:
                print(f"➕ Новые в результате:")
                for f in sorted(extra):
                    print(f"   {f}")

    print(f"\n{'='*70}")
    print(f"🏁 ДИАГНОСТИКА ЗАВЕРШЕНА")
    print(f"{'='*70}")
    print(f"\n📂 Откройте в Word файл: {output}")
    print(f"   Если рамка видна — проблема в основном коде")
    print(f"   Если рамки нет — проблема в самом шаблоне")


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("Использование: python debug_template.py рамка.docx")
        print("")
        print("Можно также указать любой .docx чтобы посмотреть структуру:")
        print("  python debug_template.py документ.docx")
        sys.exit(1)

    path = sys.argv[1]
    if not os.path.exists(path):
        print(f"❌ Файл не найден: {path}")
        sys.exit(1)

    # Структура оригинала
    dump_docx_structure(path, "ОРИГИНАЛ (рамка)")

    # Тест обработки
    test_template_processing(path)