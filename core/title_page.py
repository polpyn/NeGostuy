from docx import Document
import os

def create_title_page(data, output_path, template_path=None):
    if not template_path:
        template_path = os.path.join(os.path.dirname(__file__), 'templates', 'title_template.docx')

    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Шаблон не найден: {template_path}")

    doc = Document(template_path)

    # Собираем данные. 
    # ВНИМАНИЕ: выводим в консоль, чтобы проверить, что пришло с сайта
    print(f"DEBUG: Данные для титульника: {data}")

    # Словарь замен
    replacements = {
        "{{work_number}}": str(data.get('work_number', '')),
        "{{discipline}}": data.get('specialty_name', ''),
        "{{subject}}": data.get('subject', ''),
        "{{group_id}}": f"{data.get('group', '')}, {data.get('student_id', '')}",
        "{{student_name}}": data.get('student_name', ''),
        "{{teacher_name}}": data.get('teacher_name', ''),
        "{{city}}": data.get('city', 'Красноярск'),
        "{{year}}": data.get('year', '2026')
    }

    def process_paragraph(p):
        # Собираем весь текст параграфа
        full_text = "".join(run.text for run in p.runs)
        
        for key, value in replacements.items():
            if key in full_text:
                print(f"  --> Заменяю {key} на {value}") # Отладочное сообщение
                # Делаем замену в собранном тексте
                full_text = full_text.replace(key, value)
                
                # Сохраняем стиль (шрифт и жирность) самого первого кусочка
                if p.runs:
                    f_name = p.runs[0].font.name or "Times New Roman"
                    f_size = p.runs[0].font.size
                    f_bold = p.runs[0].bold
                else:
                    f_name, f_size, f_bold = "Times New Roman", None, None

                # Полностью перезаписываем параграф, НО сразу возвращаем стиль
                p.text = ""
                new_run = p.add_run(full_text)
                new_run.font.name = f_name
                if f_size: new_run.font.size = f_size
                new_run.bold = f_bold

    # Проходим по всем местам, где может быть текст
    for p in doc.paragraphs:
        process_paragraph(p)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    process_paragraph(p)

    doc.save(output_path)
    return output_path