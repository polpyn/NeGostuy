"""
AI-постобработка уже отформатированного DOCX.

Новый пайплайн (v2):
1) читаем абзацы из готового файла;
2) нарезаем документ на окна с полным контекстом (не только нумерованные строки);
3) отправляем окно в LLM (Gemini / OpenRouter; при явном AI_PROVIDER=ollama — локально);
4) LLM возвращает список действий в строгом JSON-формате (через responseSchema);
5) применяем безопасные правки нумерации/вложенности/жирности;
6) сохраняем документ на месте.

Поддерживаемые действия:
- bold_full    — весь абзац жирным (заголовок секции вида "1. Название")
- bold_prefix  — жирным только часть до двоеточия ("1. Как работает: ..." → "1. Как работает:" жирно)
- to_bullet    — превратить нумерованный пункт в маркированный "• ..."
- to_dash      — превратить в пункт с тире "— ..."
- renumber     — изменить номер (например "1." → "2." или "2." → "5.1.")
- keep         — ничего не менять (явный no-op)

У каждого действия есть indent_level 0..4 для поддержки глубокой вложенности.
"""

from __future__ import annotations

import json
import os
import re
import time
import urllib.error
import urllib.request
from dataclasses import dataclass
from typing import Any

from django.conf import settings
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from docx.oxml.ns import qn


NUMERIC_RE = re.compile(r"^\s*(\d+(?:\.\d+)*)\.\s+(.+)$")
BULLET_RE = re.compile(r"^\s*[—\-•]\s+(.+)$")
SUBLIST_PARENT_RE = re.compile(
    r"(преимущества|недостатки|методы|виды|этапы|включает|свойства|принципы|задачи|функции|особенности|признаки)\s*:\s*$",
    re.IGNORECASE,
)

VALID_ACTIONS = {
    "bold_full",
    "bold_prefix",
    "to_bullet",
    "to_dash",
    "renumber",
    "keep",
    "center_title",
}

# Один уровень вложенности = +1.25 см к левому полю (как базовая красная строка ГОСТ).
# Уровень 0: красная строка 1.25 см; уровни 1..n: left_indent = n × 1.25 см (без «0.75»).
INDENT_STEP_CM = 1.25
MAX_INDENT_LEVEL = 5


# ============================================================
# ИЗВЛЕЧЕНИЕ КОНТЕКСТА (окна абзацев)
# ============================================================

@dataclass
class Window:
    start: int
    end: int
    paragraphs: list[tuple[int, str]]  # (абсолютный индекс, текст)


def _collect_paragraphs(doc: Document) -> list[str]:
    return [(p.text or "").strip() for p in doc.paragraphs]


def _split_paragraphs_on_internal_newlines(doc: Document) -> int:
    """
    Абзац с несколькими строками (часто после копипаста из браузера/текста) даёт один [INDEX]
    в промпте — модель не может разметить внутри. Разбиваем на отдельные абзацы Word.
    """
    added = 0
    while True:
        done = True
        for para in list(doc.paragraphs):
            raw = (para.text or "").replace("\r\n", "\n").replace("\r", "\n")
            if "\n" not in raw:
                continue
            parts = [ln.strip() for ln in raw.split("\n") if ln.strip()]
            if len(parts) <= 1:
                continue
            para.text = parts[-1]
            for line in reversed(parts[:-1]):
                para.insert_paragraph_before(line)
            added += len(parts) - 1
            done = False
            break
        if done:
            break
    return added


def _has_list_like(text: str) -> bool:
    if not text:
        return False
    return bool(NUMERIC_RE.match(text) or BULLET_RE.match(text))


def _extract_windows(
    paragraphs: list[str],
    window_size: int = 50,
    overlap: int = 5,
    only_list_like: bool = True,
) -> list[Window]:
    """
    Нарезает документ на окна по `window_size` абзацев с перекрытием `overlap`.
    Если only_list_like=True — берём только окна с нумерацией/маркерами (быстрый путь).
    Если False — все окна (нужно для конспектов без «1. 2.»).
    """
    if not paragraphs:
        return []
    window_size = max(10, window_size)
    overlap = max(0, min(overlap, window_size - 1))
    step = window_size - overlap

    windows: list[Window] = []
    n = len(paragraphs)
    start = 0
    while start < n:
        end = min(n - 1, start + window_size - 1)
        chunk = [(i, paragraphs[i]) for i in range(start, end + 1)]
        if (not only_list_like) or any(_has_list_like(t) for _, t in chunk):
            windows.append(Window(start=start, end=end, paragraphs=chunk))
        if end >= n - 1:
            break
        start += step
    return windows


# ============================================================
# DEBUG LOG
# ============================================================

def _write_debug_log(
    payload: dict[str, Any],
    *,
    run_ts: str | None = None,
) -> str | None:
    enabled = str(getattr(settings, "AI_POSTPROCESS_DEBUG", "0")).lower() in {
        "1", "true", "yes", "on"
    }
    if not enabled:
        return None
    debug_dir = getattr(settings, "AI_POSTPROCESS_DEBUG_DIR", "").strip()
    if not debug_dir:
        debug_dir = os.path.join(str(settings.BASE_DIR), "debug_ai")
    os.makedirs(debug_dir, exist_ok=True)
    ts = run_ts or time.strftime("%Y%m%d_%H%M%S")
    out_path = os.path.join(debug_dir, f"ai_postprocess_{ts}.json")
    try:
        with open(out_path, "w", encoding="utf-8") as f:
            json.dump(payload, f, ensure_ascii=False, indent=2)
        return out_path
    except OSError:
        return None


def _trace_enabled() -> bool:
    return str(getattr(settings, "AI_POSTPROCESS_TRACE", "0")).lower() in {
        "1", "true", "yes", "on",
    }


def _trace_files_enabled() -> bool:
    return str(getattr(settings, "AI_POSTPROCESS_TRACE_FILES", "0")).lower() in {
        "1", "true", "yes", "on",
    }


def _trace_console_enabled() -> bool:
    if not _trace_enabled():
        return False
    return str(getattr(settings, "AI_POSTPROCESS_TRACE_CONSOLE", "1")).lower() in {
        "1", "true", "yes", "on",
    }


def _build_window_llm_trace(
    window_index: int,
    w: Window,
    prompt: str,
    actions: list[dict[str, Any]],
    err: str,
    meta: dict[str, Any],
) -> dict[str, Any]:
    """Полная трассировка одного окна: что ушло в LLM и что вернулось (для отладки «первых страниц»)."""
    head_n = int(getattr(settings, "AI_POSTPROCESS_TRACE_PROMPT_HEAD", 8000) or 8000)
    tail_n = int(getattr(settings, "AI_POSTPROCESS_TRACE_PROMPT_TAIL", 4000) or 4000)
    rmax = int(getattr(settings, "AI_POSTPROCESS_TRACE_RESPONSE_MAX", 200000) or 200000)

    inner = (meta or {}).get("llm_inner_text_full") or ""
    if not inner:
        inner = (meta or {}).get("model_text_preview") or ""

    inner_len = len(inner)
    if inner_len > rmax:
        inner_stored = inner[:rmax] + f"\n... [обрезано при записи в JSON, всего {inner_len} символов]"
    else:
        inner_stored = inner

    prompt_len = len(prompt)
    prompt_head = prompt[:head_n]
    prompt_tail = prompt[-tail_n:] if prompt_len > head_n else ""

    meta_light: dict[str, Any] = {}
    for k, v in (meta or {}).items():
        if k in ("llm_inner_text_full", "model_text_preview", "attempts"):
            continue
        if k == "gemini_http_attempts" and isinstance(v, list):
            meta_light[k] = v[:12]
        else:
            try:
                json.dumps(v, default=str)
                meta_light[k] = v
            except TypeError:
                meta_light[k] = str(v)[:500]

    trace: dict[str, Any] = {
        "window_index": window_index,
        "paragraph_index_range": [w.start, w.end],
        "paragraph_count_in_window": len(w.paragraphs),
        "hint": (
            "Начало документа обычно в окне 0. Если абзацы 0–1 без действий в ответе — "
            "модель могла считать их вводом без структуры; абзацы в таблицах не входят в этот список."
        ),
        "paragraphs_in_window_preview": [
            {"index": idx, "text_head": (txt or "")[:300]}
            for idx, txt in w.paragraphs[:35]
        ],
        "prompt_total_chars": prompt_len,
        "prompt_head": prompt_head,
        "prompt_tail": prompt_tail,
        "llm_inner_response_total_chars": inner_len,
        "llm_inner_response_text": inner_stored,
        "parsed_actions_count": len(actions),
        "parsed_actions": actions,
        "error": err or None,
        "provider_meta": meta_light,
    }
    att = (meta or {}).get("attempts")
    if att:
        trace["provider_fallback_attempts"] = att
    return trace


def _write_window_trace_files(
    debug_dir: str,
    run_ts: str,
    window_index: int,
    prompt: str,
    inner_text: str,
    err: str,
) -> None:
    sub = os.path.join(debug_dir, f"llm_trace_{run_ts}")
    os.makedirs(sub, exist_ok=True)
    base = os.path.join(sub, f"w{window_index}")
    try:
        with open(base + "_prompt.txt", "w", encoding="utf-8") as f:
            f.write(prompt)
        with open(base + "_response.txt", "w", encoding="utf-8") as f:
            f.write(inner_text or err or "")
    except OSError:
        pass


def _windows_digest(
    windows: list[Window],
    prompts_and_results: list[tuple[str, list, str, dict]],
) -> list[dict[str, Any]]:
    """Короткая сводка без полных текстов — всегда при AI_POSTPROCESS_DEBUG."""
    out = []
    for i, w in enumerate(windows):
        if i >= len(prompts_and_results):
            break
        prompt, actions, err, meta = prompts_and_results[i]
        inner = (meta or {}).get("llm_inner_text_full") or ""
        out.append({
            "window_index": i,
            "range": [w.start, w.end],
            "prompt_chars": len(prompt),
            "inner_response_chars": len(inner),
            "actions_returned": len(actions),
            "has_error": bool(err),
            "error_head": (err or "")[:240],
            "first_paragraph_index_in_window": w.paragraphs[0][0] if w.paragraphs else None,
            "first_paragraph_preview": (w.paragraphs[0][1] or "")[:160] if w.paragraphs else "",
        })
    return out


# ============================================================
# ПРОМПТ
# ============================================================

_SYSTEM_INSTRUCTION = (
    "Ты редактор структуры учебных отчётов на русском языке. "
    "Твоя задача — анализировать абзацы документа и возвращать список действий "
    "для приведения нумерации, вложенности, выравнивания и жирности к аккуратному учебному стилю."
)

_PROMPT_RULES = """Тебе дают окно абзацев из документа в формате [INDEX] текст.
Ты возвращаешь JSON со списком действий: что сделать с каждым абзацем, который нужно изменить.

Доступные действия (поле "action"):
- bold_full     — сделать весь абзац жирным (заголовок секции). Пример: "1. Файловый сервер (FS)".
- bold_prefix   — жирным только часть ДО двоеточия. Пример: "1. Как работает: ..." → жирным "1. Как работает:".
- center_title  — заголовок по центру страницы, весь абзац жирным (без смены текста). Для строк про практическую работу и «ход работы» (см. правило ниже).
- to_bullet     — маркер "• ...": для нумерованных пунктов или (см. правило 10) для коротких немаркированных строк перечня под заголовком при indent_level >= 1.
- to_dash       — то же с тире "— ...".
- renumber      — изменить номер. Укажи в new_number новый номер вида "5." или "5.1.".
- keep          — не менять (не обязателен, просто пропусти абзац).

Каждое действие ДОЛЖНО содержать поле indent_level (0..4):
- 0 = верхний уровень: красная строка 1.25 см (обычный абзац, заголовок «N. …», и этикетки «Суть:», «Главная задача:» с пояснением — НЕ пункт списка).
- 1 = левый отступ 1.25 см (редко — мелкий подуровень без номера секции).
- 2 = левый отступ 2.5 см — только для to_bullet / to_dash (настоящие пункты перечня под «N. …», строки с «•»).
- 3..4 = глубже (+1.25 см за уровень).

СТРОГИЕ ПРАВИЛА:

1. ЗАГОЛОВКИ СЕКЦИЙ. Короткие абзацы вида "N. Название" БЕЗ двоеточия в содержимом и без глагольных хвостов
   ("Файловый сервер (FS)", "Системы с разделением времени", "Коммутация пакетов (КП)") — это заголовки.
   То же для **без номера**: отдельная строка-название блока — «Системы … (СРВ / …)», «Системы «Терминал — Хост»»,
   «Системы «Клиент — Сервер»», строка с **ключевой дилеммой/тезисом, заканчивающаяся двоеточием** — это заголовки уровня секции.
   Действие: bold_full (или bold_prefix, если только часть до «:» должна быть жирной), indent_level = уровень секции.
   Если после такого заголовка идут подпункты вида "1. Prefix:", "2. Prefix:", "3. Prefix:" —
   это СТОПРОЦЕНТНО заголовок (bold_full), а не обычный пункт.

2. ПОДПУНКТЫ С ПОЯСНЕНИЕМ. Абзацы вида "N. Prefix: пояснение..." ПОД заголовком секции —
   это элементы вложенного списка. Действие: to_bullet, indent_level = 2 (отступ 2.5 см под «N. …»).
   Программа сама оставит часть до ":" жирной внутри буллета.
   ИСКЛЮЧЕНИЕ — раздел «ХОД РАБОТЫ» (см. правило 9): там такие строки НЕ переводить в буллеты.

3. ВАЖНО про to_bullet и to_dash. Обычно — к абзацам с "N." или маркером "—/-/•".
   Для "Суть: ...", "Главная задача ...: ..." — всегда bold_prefix, НЕ to_bullet: это поясняющий текст после заголовка «N. …»,
   оформляется как обычный абзац (indent_level = 0, красная строка 1.25 см), а не как вложенный пункт на 2.5 см.
   Исключение: несколько коротких однотипных строк под заголовком (например «Терминалом, когда …», «Хостом, когда …») —
   to_bullet с indent_level 2 под строкой «N. …». Не ставь to_bullet длинным монолитным абзацам.

4. ПОДЗАГОЛОВКИ-ЭТИКЕТКИ. Абзац вида "Суть:", "Главная задача:", "Важное наблюдение:",
   "Преимущества:", "Недостатки:" (метка с двоеточием и пояснением) — подзаголовок-этикетка внутри секции.
   Действие: bold_prefix, indent_level = 0 всегда (в т.ч. сразу под «N. Название секции»). На 2.5 см сдвигаются только буллеты/to_bullet.

5. СЛОМАННАЯ НУМЕРАЦИЯ. Если видишь серию "1. X", "1. Y", "1. Z", "1. W" (нумерация обнулилась на том же уровне)
   и это должен быть плоский список — renumber с последовательными номерами 1, 2, 3, 4, 5...

6. ТЕМАТИЧЕСКАЯ ВЛОЖЕННОСТЬ. Если по смыслу пункт относится как подпункт к предыдущему
   (например под "2. Расчёт метрики" идёт пункт, который по сути — часть расчёта),
   используй renumber с new_number вида "2.1.".

7. РИСУНКИ И ТАБЛИЦЫ. НЕ ТРОГАЙ абзацы, начинающиеся с "Рисунок " или "Таблица ".

8. ОБЫЧНЫЙ ТЕКСТ. Если абзац — обычный прозаический текст без нумерации, без двоеточия-подзаголовка
   и не похож на заголовок — НЕ включай его в actions. Пропусти его. Исключения — правило 10 и абзац ниже.
   Под абзацем-заголовком раздела «N. …» (например «1. Системы с разделением времени»): отдельные короткие абзацы
   с требованиями, ограничениями или однотипными утверждениями об одном объекте
   («Пользователи не должны…», «Данные одного пользователя должны…», «Нужно дать инструменты…»)
   — это перечень под разделом, не «проза для пропуска»: оформляй to_bullet, indent_level 2 (под «N. …»).

9. ПРАКТИЧЕСКАЯ РАБОТА И ХОД РАБОТЫ.
   - Заголовки-титулы: «Практическая работа», «Практическая работа №…», «Лабораторная работа…»,
     варианты с «по дисциплине», отдельная строка «Ход работы» (любое нормальное написание, с двоеточием или без)
     — всегда center_title, indent_level 0. Не превращай их в to_bullet / to_dash.
   - После «Ход работы» до следующего крупного раздела строки «1. …», «2. …», «3. …» — последовательные шаги отчёта
     (файлы, модели, формы, представления, скриншоты). Для них ЗАПРЕЩЕНЫ to_bullet и to_dash: остаётся нумерация 1. 2. 3.
     При сбое нумерации — только renumber; если уже верно — keep или не включай в actions.
     bold_prefix только при явном формате «N. Краткий заголовок: продолжение текста»; длинные шаги без двоеточия не трогай (keep).


10. ДОКУМЕНТ БЕЗ НУМЕРАЦИИ. Даже без «1. 2. 3.» выделяй структуру по смыслу: названия блоков — bold_full;
   этикетки «Суть:», «Главная задача …:» — bold_prefix, indent_level 0; однородные короткие пункты-перечисления — to_bullet (indent 2 под «N. …», иначе 1).
   Если в окне есть разделы «1. …», «2. …», подпункты из п.8 — to_bullet с indent_level 2.
   Смотри соседние абзацы в окне. Начало документа (вводный абзац, тезис с двоеточием, первые блоки) обрабатывай так же приоритетно, как середину.

ВОЗВРАЩАЙ ТОЛЬКО те абзацы, которые реально нужно изменить.
Если менять ничего не нужно — верни пустой список {"actions": []}.

Пример желаемого результата для фрагмента:
  [35] Эволюция четырёх архитектур
  [37] 1. Файловый сервер (FS)
  [38] 1. Как работает: Клиентская программа...
  [39] 2. Недостатки: Дикий трафик...
  [40] 3. Вывод: Это примитивный...
  [41] 2. Доступ к удалённым данным (RDA)
  [42] 1. Как работает: Теперь на сервере...

Ответ:
{
  "actions": [
    {"index": 37, "action": "bold_full",  "indent_level": 0, "reason": "заголовок секции"},
    {"index": 38, "action": "to_bullet",  "indent_level": 2, "reason": "подпункт под 'Файловый сервер'"},
    {"index": 39, "action": "to_bullet",  "indent_level": 2},
    {"index": 40, "action": "to_bullet",  "indent_level": 2},
    {"index": 41, "action": "bold_full",  "indent_level": 0, "reason": "заголовок секции"},
    {"index": 42, "action": "to_bullet",  "indent_level": 2}
  ]
}
"""


def _build_window_prompt(window: Window) -> str:
    lines = [f"[{idx}] {text}" if text else f"[{idx}] <пустая строка>"
             for idx, text in window.paragraphs]
    window_block = "\n".join(lines)
    return (
        _SYSTEM_INSTRUCTION + "\n\n" +
        _PROMPT_RULES + "\n\n" +
        f"Окно абзацев (индексы {window.start}..{window.end}):\n" +
        window_block + "\n\n" +
        'Верни ТОЛЬКО JSON вида {"actions":[...]} без пояснений и markdown.'
    )


# ============================================================
# JSON SCHEMA ДЛЯ GEMINI
# ============================================================

_GEMINI_RESPONSE_SCHEMA = {
    "type": "OBJECT",
    "properties": {
        "actions": {
            "type": "ARRAY",
            "items": {
                "type": "OBJECT",
                "properties": {
                    "index": {"type": "INTEGER"},
                    "action": {
                        "type": "STRING",
                        "enum": [
                            "bold_full",
                            "bold_prefix",
                            "center_title",
                            "to_bullet",
                            "to_dash",
                            "renumber",
                            "keep",
                        ],
                    },
                    "indent_level": {"type": "INTEGER"},
                    "new_number": {"type": "STRING"},
                    "reason": {"type": "STRING"},
                },
                "required": ["index", "action"],
            },
        }
    },
    "required": ["actions"],
}


# ============================================================
# HTTP УТИЛИТЫ (с поддержкой прокси)
# ============================================================

def _make_opener(proxy_url: str = ""):
    """
    Создаёт urllib opener. Если задан proxy_url — все запросы пойдут через него.
    Иначе использует системные HTTP_PROXY / HTTPS_PROXY (или прямое соединение).
    """
    handlers = []
    if proxy_url:
        handlers.append(urllib.request.ProxyHandler({
            "http": proxy_url,
            "https": proxy_url,
        }))
        return urllib.request.build_opener(*handlers)
    return urllib.request.build_opener()


def _gemini_quota_exhausted_error(err: str) -> bool:
    """True, если 429 из-за дневной/проектной квоты — повторы только удлиняют ожидание."""
    if not err:
        return False
    lower = err.lower()
    return (
        "quota exceeded" in lower
        or "resource_exhausted" in lower
        or "free_tier" in lower
    )


def _http_post_json(
    url: str,
    body: dict[str, Any],
    headers: dict[str, str],
    timeout: int,
    proxy_url: str = "",
) -> tuple[str, str, int]:
    """
    Возвращает (response_text, error, http_status).
    При успехе: error == "", status=200.
    """
    req = urllib.request.Request(
        url,
        data=json.dumps(body).encode("utf-8"),
        headers={"Content-Type": "application/json", **headers},
        method="POST",
    )
    opener = _make_opener(proxy_url)
    try:
        with opener.open(req, timeout=timeout) as resp:
            return resp.read().decode("utf-8"), "", resp.status
    except urllib.error.HTTPError as exc:
        try:
            err_body = exc.read().decode("utf-8", errors="replace")[:2000]
        except Exception:
            err_body = ""
        return "", f"http {exc.code}: {err_body}", exc.code
    except (urllib.error.URLError, TimeoutError, ValueError) as exc:
        return "", f"network error: {exc}", 0


# ============================================================
# GEMINI API (напрямую через Google AI Studio)
# ============================================================

def _call_gemini(prompt: str) -> tuple[list[dict[str, Any]], str, dict[str, Any]]:
    api_key = (getattr(settings, "GEMINI_API_KEY", "") or "").strip()
    if not api_key:
        return [], "GEMINI_API_KEY is empty", {}

    model = (getattr(settings, "GEMINI_MODEL", "gemini-2.5-flash") or "").strip()
    base_url = (getattr(settings, "GEMINI_BASE_URL",
                        "https://generativelanguage.googleapis.com/v1beta")).rstrip("/")
    timeout = int(getattr(settings, "GEMINI_TIMEOUT", 120))
    proxy = (getattr(settings, "GEMINI_PROXY", "") or "").strip()

    url = f"{base_url}/models/{model}:generateContent?key={api_key}"
    body = {
        "contents": [{"role": "user", "parts": [{"text": prompt}]}],
        "generationConfig": {
            "temperature": 0.1,
            "responseMimeType": "application/json",
            "responseSchema": _GEMINI_RESPONSE_SCHEMA,
        },
    }
    max_retries = int(getattr(settings, "GEMINI_MAX_RETRIES", 5))
    backoff = float(getattr(settings, "GEMINI_RETRY_BACKOFF_SEC", 2.0))
    max_wait = float(getattr(settings, "GEMINI_RETRY_MAX_WAIT_SEC", 60.0))
    #503 «high demand» у бесплатного API часто лечится повтором; иначе первое окно (начало документа) остаётся без правок.
    transient = {429, 500, 502, 503, 504}

    t_start = time.time()
    raw, err, status = "", "no request", 0
    http_attempts: list[dict[str, Any]] = []
    for attempt in range(max_retries + 1):
        t0 = time.time()
        raw, err, status = _http_post_json(url, body, {}, timeout, proxy_url=proxy)
        http_attempts.append({
            "n": attempt,
            "request_ms": int((time.time() - t0) * 1000),
            "status": status,
            "error": err or None,
        })
        if not err:
            break
        if _gemini_quota_exhausted_error(err):
            break
        # Таймаут/обрыв (WinError 10060 и т.п., status 0) — повторяем; иначе LLM молчит и остаётся слабая эвристика.
        err_l = (err or "").lower()
        networkish = status == 0 and (
            "network error" in err_l or "10060" in err or "timed out" in err_l or "urlopen" in err_l
        )
        if networkish and attempt < max_retries:
            time.sleep(min(backoff * (2**attempt), max_wait))
            continue
        if status not in transient or attempt >= max_retries:
            break
        time.sleep(min(backoff * (2**attempt), max_wait))

    elapsed_ms = int((time.time() - t_start) * 1000)
    meta: dict[str, Any] = {
        "request_ms": elapsed_ms,
        "model": model,
        "status": status,
        "gemini_http_attempts": http_attempts,
    }
    if err:
        meta["error_code"] = status
        return [], f"gemini {err}", meta

    try:
        data = json.loads(raw)
    except Exception as exc:
        meta["raw_preview"] = raw[:2000]
        return [], f"gemini invalid outer json: {exc}", meta

    candidates = data.get("candidates") or []
    if not candidates:
        meta["raw_preview"] = raw[:2000]
        return [], "gemini no candidates", meta

    parts = ((candidates[0].get("content") or {}).get("parts") or [])
    text = "".join(p.get("text", "") for p in parts if isinstance(p, dict))
    meta["model_text_preview"] = text[:4000]
    meta["llm_inner_text_full"] = text

    if not text.strip():
        return [], "gemini empty response", meta

    try:
        parsed = json.loads(text)
    except Exception as exc:
        meta["inner_json_parse_error"] = str(exc)
        return [], f"gemini invalid inner json: {exc}", meta

    actions = parsed.get("actions") if isinstance(parsed, dict) else None
    if not isinstance(actions, list):
        return [], "gemini no actions array", meta
    return [a for a in actions if isinstance(a, dict)], "", meta


# ============================================================
# OPENROUTER API (OpenAI-совместимый, работает из РФ без VPN)
# ============================================================

def _call_openrouter(prompt: str) -> tuple[list[dict[str, Any]], str, dict[str, Any]]:
    api_key = (getattr(settings, "OPENROUTER_API_KEY", "") or "").strip()
    if not api_key:
        return [], "OPENROUTER_API_KEY is empty", {}

    model = (getattr(settings, "OPENROUTER_MODEL", "google/gemini-2.5-flash") or "").strip()
    base_url = (getattr(settings, "OPENROUTER_BASE_URL",
                        "https://openrouter.ai/api/v1")).rstrip("/")
    timeout = int(getattr(settings, "OPENROUTER_TIMEOUT", 120))

    url = f"{base_url}/chat/completions"
    body = {
        "model": model,
        "messages": [
            {"role": "system", "content": _SYSTEM_INSTRUCTION},
            {"role": "user", "content": prompt},
        ],
        "temperature": 0.1,
        "response_format": {"type": "json_object"},
    }
    headers = {
        "Authorization": f"Bearer {api_key}",
        "HTTP-Referer": "https://negostuy.local",
        "X-Title": "NeGostuy",
    }
    t_start = time.time()
    raw, err, status = _http_post_json(url, body, headers, timeout)
    elapsed_ms = int((time.time() - t_start) * 1000)
    meta: dict[str, Any] = {"request_ms": elapsed_ms, "model": model, "status": status}
    if err:
        meta["error_code"] = status
        return [], f"openrouter {err}", meta

    try:
        data = json.loads(raw)
    except Exception as exc:
        meta["raw_preview"] = raw[:2000]
        return [], f"openrouter invalid outer json: {exc}", meta

    choices = data.get("choices") or []
    if not choices:
        meta["raw_preview"] = raw[:2000]
        return [], "openrouter no choices", meta
    msg = (choices[0].get("message") or {}).get("content", "")
    meta["model_text_preview"] = msg[:4000]
    meta["llm_inner_text_full"] = msg
    if not msg.strip():
        return [], "openrouter empty content", meta

    actions = _extract_actions_from_text(msg)
    if not actions:
        return [], "openrouter parse empty", meta
    return actions, "", meta


# ============================================================
# OLLAMA API (fallback)
# ============================================================

def _extract_actions_from_text(text: str) -> list[dict[str, Any]]:
    raw = (text or "").strip()
    if not raw:
        return []
    raw = re.sub(r"^```(?:json)?\s*", "", raw, flags=re.IGNORECASE)
    raw = re.sub(r"\s*```$", "", raw)
    try:
        parsed = json.loads(raw)
    except Exception:
        parsed = None

    if parsed is None:
        start = raw.find("{")
        end = raw.rfind("}")
        if start >= 0 and end > start:
            try:
                parsed = json.loads(raw[start:end + 1])
            except Exception:
                parsed = None

    if isinstance(parsed, list):
        return [x for x in parsed if isinstance(x, dict)]
    if isinstance(parsed, dict):
        for key in ("actions", "recommendations", "items", "result"):
            val = parsed.get(key)
            if isinstance(val, list):
                return [x for x in val if isinstance(x, dict)]
        data = parsed.get("data")
        if isinstance(data, list):
            return [x for x in data if isinstance(x, dict)]
        if isinstance(data, dict):
            items = data.get("items") or data.get("actions")
            if isinstance(items, list):
                return [x for x in items if isinstance(x, dict)]
    return []


def _call_ollama(prompt: str) -> tuple[list[dict[str, Any]], str, dict[str, Any]]:
    base_url = getattr(settings, "OLLAMA_BASE_URL", "http://127.0.0.1:11434")
    model = (getattr(settings, "OLLAMA_MODEL", "") or "").strip()
    timeout = int(getattr(settings, "OLLAMA_TIMEOUT", 240))
    if not model:
        return [], "OLLAMA_MODEL is empty", {}

    url = f"{base_url.rstrip('/')}/api/generate"
    body = {
        "model": model,
        "prompt": prompt,
        "stream": False,
        "options": {"temperature": 0.1},
        "format": "json",
    }
    req = urllib.request.Request(
        url,
        data=json.dumps(body).encode("utf-8"),
        headers={"Content-Type": "application/json"},
        method="POST",
    )
    t_start = time.time()
    try:
        with urllib.request.urlopen(req, timeout=timeout) as resp:
            raw = resp.read().decode("utf-8")
    except (urllib.error.URLError, TimeoutError, ValueError) as exc:
        return [], f"ollama error: {exc}", {}
    elapsed_ms = int((time.time() - t_start) * 1000)
    meta = {"request_ms": elapsed_ms, "model": model}
    try:
        data = json.loads(raw)
    except Exception as exc:
        meta["raw_preview"] = raw[:2000]
        return [], f"ollama outer json error: {exc}", meta
    model_text = (data or {}).get("response", "")
    meta["model_text_preview"] = model_text[:4000]
    meta["llm_inner_text_full"] = model_text
    actions = _extract_actions_from_text(model_text)
    return actions, "", meta


# ============================================================
# ДИСПЕТЧЕР LLM
# ============================================================

def _resolve_provider() -> str:
    """
    Возвращает первичный провайдер. При auto приоритет:
    openrouter (если ключ задан) → gemini (если ключ задан) → ollama.
    Цепочка fallback в _fallback_chain не включает Ollama — только OpenRouter и Gemini.
    """
    provider = (getattr(settings, "AI_PROVIDER", "auto") or "auto").strip().lower()
    if provider in {"gemini", "ollama", "openrouter"}:
        return provider
    if (getattr(settings, "OPENROUTER_API_KEY", "") or "").strip():
        return "openrouter"
    if (getattr(settings, "GEMINI_API_KEY", "") or "").strip():
        return "gemini"
    return "ollama"


_PROVIDER_CALLS = {
    "gemini": _call_gemini,
    "openrouter": _call_openrouter,
    "ollama": _call_ollama,
}


def _fallback_chain(primary: str) -> list[str]:
    """Порядок попыток: сначала primary, потом OpenRouter/Gemini при наличии ключей (без Ollama)."""
    chain = [primary]
    for candidate in ("openrouter", "gemini"):
        if candidate == primary:
            continue
        if candidate == "gemini" and not (getattr(settings, "GEMINI_API_KEY", "") or "").strip():
            continue
        if candidate == "openrouter" and not (getattr(settings, "OPENROUTER_API_KEY", "") or "").strip():
            continue
        chain.append(candidate)
    return chain


def _print_llm_route_banner() -> None:
    """
    Печатает фактический первичный провайдер и имя модели.
    При AI_PROVIDER=auto и заданном OPENROUTER_API_KEY первым идёт OpenRouter — тогда
    счётчик RPD в Google AI Studio для GEMINI_MODEL не меняется (запросы не в Google напрямую).
    """
    primary = _resolve_provider()
    gemini_m = (getattr(settings, "GEMINI_MODEL", "") or "").strip() or "gemini-2.5-flash"
    or_m = (getattr(settings, "OPENROUTER_MODEL", "") or "").strip() or "google/gemini-2.5-flash"
    has_gem_key = bool((getattr(settings, "GEMINI_API_KEY", "") or "").strip())
    has_or_key = bool((getattr(settings, "OPENROUTER_API_KEY", "") or "").strip())
    explicit = (getattr(settings, "AI_PROVIDER", "auto") or "auto").strip().lower()

    if primary == "gemini":
        print(f"  [LLM] Запросы: Google Gemini API напрямую, модель `{gemini_m}`.")
    elif primary == "openrouter":
        print(f"  [LLM] Запросы: OpenRouter (первичный провайдер), модель `{or_m}`.")
        if has_gem_key and explicit == "auto":
            print(
                "  [LLM] Внимание: при AI_PROVIDER=auto первым выбран OpenRouter — "
                "квота в Google AI Studio для GEMINI_MODEL не расходуется. "
                "Нужен только Google и переменная GEMINI_MODEL — поставьте AI_PROVIDER=gemini "
                "(или уберите OPENROUTER_API_KEY)."
            )
    else:
        ollama_m = (getattr(settings, "OLLAMA_MODEL", "") or "").strip() or "(OLLAMA_MODEL пуст)"
        print(f"  [LLM] Запросы: Ollama, модель `{ollama_m}`.")
    if not has_or_key and not has_gem_key:
        print("  [LLM] Нет ни OPENROUTER_API_KEY, ни GEMINI_API_KEY — LLM не сможет ответить.")


def _call_llm(prompt: str) -> tuple[list[dict[str, Any]], str, dict[str, Any]]:
    primary = _resolve_provider()
    last_err = ""
    last_meta: dict[str, Any] = {}
    attempts: list[dict[str, Any]] = []
    for prov in _fallback_chain(primary):
        call = _PROVIDER_CALLS.get(prov)
        if call is None:
            continue
        actions, err, meta = call(prompt)
        meta = dict(meta)
        meta["provider"] = prov
        if err:
            meta["error"] = err
        attempts.append(dict(meta))
        meta["attempts"] = attempts
        if actions:
            return actions, "", meta
        if not err:
            # Валидный пустой ответ — не пробуем следующего провайдера.
            return actions, err, meta
        last_err = err
        last_meta = meta
    if last_meta:
        out = dict(last_meta)
        out["attempts"] = attempts
        return [], last_err, out
    return [], last_err, {"attempts": attempts}


# ============================================================
# ПРИМЕНЕНИЕ ДЕЙСТВИЙ К DOCX
# ============================================================

def _set_run_font(run, bold: bool = False) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    run.font.bold = bold
    try:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    except Exception:
        pass


def _replace_runs(para, segments: list[tuple[str, bool]]) -> None:
    for r in list(para.runs):
        try:
            para._element.remove(r._element)
        except Exception:
            pass
    for text, bold in segments:
        if not text:
            continue
        run = para.add_run(text)
        _set_run_font(run, bold=bold)


def _apply_structure_style(para, indent_level: int = 0) -> None:
    """Выставляет выравнивание, межстрочный интервал и отступы для любого уровня 0..4."""
    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.5
    level = max(0, min(int(indent_level or 0), MAX_INDENT_LEVEL))
    if level == 0:
        para.paragraph_format.left_indent = Cm(0)
        para.paragraph_format.first_line_indent = Cm(INDENT_STEP_CM)
    else:
        para.paragraph_format.left_indent = Cm(INDENT_STEP_CM * level)
        para.paragraph_format.first_line_indent = Cm(0)


def _normalize_action(item: dict[str, Any]) -> dict[str, Any] | None:
    """Приводит action к каноническому виду, с запасом по неточным ответам моделей."""
    try:
        idx = int(item.get("index"))
    except Exception:
        return None

    raw_kind = str(item.get("action", "")).strip().lower()
    aliases = {
        "bold": "bold_full",
        "make_bold": "bold_full",
        "title_bold": "bold_full",
        "bullet": "to_bullet",
        "bulleted": "to_bullet",
        "to-bullet": "to_bullet",
        "dash": "to_dash",
        "to-dash": "to_dash",
        "bold_until_colon": "bold_prefix",
        "prefix_bold": "bold_prefix",
        "subitem": "renumber",
        "center": "center_title",
        "title_center": "center_title",
        "centered_title": "center_title",
    }
    kind = aliases.get(raw_kind, raw_kind)
    if kind not in VALID_ACTIONS:
        return None

    try:
        indent_level = int(item.get("indent_level", 0) or 0)
    except Exception:
        indent_level = 0
    indent_level = max(0, min(indent_level, MAX_INDENT_LEVEL))

    return {
        "index": idx,
        "action": kind,
        "indent_level": indent_level,
        "new_number": str(item.get("new_number", "")).strip(),
        "new_text": str(item.get("new_text", "")).strip(),
        "reason": str(item.get("reason", "")).strip(),
    }


def _is_protected(text: str) -> bool:
    low = (text or "").strip().lower()
    return low.startswith("рисунок ") or low.startswith("таблица ")


def _apply_single_action(doc: Document, action: dict[str, Any]) -> dict[str, Any] | None:
    idx = action["index"]
    if idx < 0 or idx >= len(doc.paragraphs):
        return None
    para = doc.paragraphs[idx]
    old_text = (para.text or "").strip()
    if not old_text or _is_protected(old_text):
        return None

    kind = action["action"]
    level = action["indent_level"]

    if kind == "keep":
        return None

    if kind == "center_title":
        if len(old_text) > 300:
            return None
        _replace_runs(para, [(old_text, True)])
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.left_indent = Cm(0)
        para.paragraph_format.first_line_indent = Cm(0)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = 1.5
        return {"index": idx, "action": kind, "indent_level": 0, "old_text": old_text}

    if kind == "bold_full":
        mnum = NUMERIC_RE.match(old_text)
        is_numbered = bool(mnum)
        content = mnum.group(2).strip() if mnum else old_text
        lowered = content.lower()
        ends_with_period = content.rstrip().endswith(".")
        has_colon = ":" in content
        has_verbs = bool(re.search(
            r"\b(должен|должны|может|могут|нужно|нужны|является|это было|являются|"
            r"представля|означает|состоит|зависит|позволя|обеспечива|использу)\b",
            lowered,
        ))

        # Не жирним целиком длинные прозаические абзацы — это гарантированно ошибка
        # (такие куски обычно попадают сюда, когда модель перепутала прозу с заголовком).
        if len(content) > 150 and not is_numbered:
            return None
        if len(content) > 200:
            return None
        # Длинный нумерованный пункт-с-пояснением: жирним только заголовочную часть
        # до первого ". " или ":". Если таких разделителей нет — отказываемся
        # (не делаем жирным полный абзац прозы).
        if is_numbered and len(content) > 90:
            num = mnum.group(1)
            cut = -1
            for p in (content.find(". "), content.find(":")):
                if p > 0 and (cut < 0 or p < cut):
                    cut = p
            if cut <= 0:
                return None
            is_colon = (content.find(":") == cut)
            head = content[:cut + (0 if is_colon else 1)].strip()
            tail = content[cut + (1 if is_colon else 2):].strip()
            _replace_runs(
                para,
                [(f"{num}. {head}", True), (f" {tail}" if tail else "", False)],
            )
            _apply_structure_style(para, indent_level=level)
            return {"index": idx, "action": kind, "indent_level": level, "old_text": old_text}

        # Короткий заголовок (нумерованный или нет). Запрещаем глагольные хвосты,
        # которые почти всегда признак обычного предложения, а не заголовка.
        if has_verbs and ends_with_period:
            return None
        # Если это короткая строка-этикетка с двоеточием — предпочитаем bold_prefix,
        # чтобы жирным был только префикс, а не «хвост».
        if has_colon and not is_numbered:
            head_text, tail_text = old_text.split(":", 1)
            _replace_runs(
                para,
                [
                    (head_text.strip() + ":", True),
                    (" " + tail_text.lstrip() if tail_text.strip() else "", False),
                ],
            )
            _apply_structure_style(para, indent_level=level)
            return {
                "index": idx,
                "action": "bold_prefix",
                "indent_level": level,
                "old_text": old_text,
                "fallback_from": "bold_full",
            }
        _replace_runs(para, [(old_text, True)])
        _apply_structure_style(para, indent_level=level)
        return {"index": idx, "action": kind, "indent_level": level, "old_text": old_text}

    if kind == "bold_prefix":
        mnum = NUMERIC_RE.match(old_text)
        if mnum and ":" in mnum.group(2):
            num, content = mnum.group(1), mnum.group(2).strip()
            head, tail = content.split(":", 1)
            _replace_runs(
                para,
                [(f"{num}. ", False), (head.strip() + ":", True), (" " + tail.lstrip(), False)],
            )
            _apply_structure_style(para, indent_level=level)
            return {"index": idx, "action": kind, "indent_level": level, "old_text": old_text}
        # «Суть: …», «Главная задача …: …» без ведущего «N.»
        if ":" in old_text:
            head_text, tail_text = old_text.split(":", 1)
            head = head_text.strip()
            tail = tail_text.strip()
            if head and len(head) <= 120 and len(head.split()) <= 14:
                _replace_runs(
                    para,
                    [(head + ":", True), (" " + tail if tail else "", False)],
                )
                _apply_structure_style(para, indent_level=level)
                return {"index": idx, "action": kind, "indent_level": level, "old_text": old_text}
        return None

    if kind in {"to_bullet", "to_dash"}:
        marker = "•" if kind == "to_bullet" else "—"
        m_num = NUMERIC_RE.match(old_text)
        m_bul = BULLET_RE.match(old_text)
        core: str | None = None
        if m_num:
            core = m_num.group(2).strip()
        elif m_bul:
            core = m_bul.group(1).strip()
        else:
            stripped = old_text.strip()
            lev = int(level or 0)
            if 0 < len(stripped) <= 400 and (
                lev >= 1 or (lev == 0 and len(stripped) <= 280)
            ):
                core = stripped
        if not core:
            # Обычный прозаический абзац без нумерации — превращать в буллет
            # опасно (часто ломает повествование). Откатываемся на bold_prefix,
            # если есть короткий префикс с двоеточием. Иначе пропускаем.
            if ":" in old_text:
                head = old_text.split(":", 1)[0].strip()
                if 0 < len(head) <= 80 and len(head.split()) <= 10:
                    head_text, tail_text = old_text.split(":", 1)
                    _replace_runs(
                        para,
                        [
                            (head_text.strip() + ":", True),
                            (" " + tail_text.lstrip(), False),
                        ],
                    )
                    _apply_structure_style(para, indent_level=level)
                    return {
                        "index": idx,
                        "action": "bold_prefix",
                        "indent_level": level,
                        "old_text": old_text,
                        "fallback_from": kind,
                    }
            return None
        # Если остался префикс с двоеточием — оставим его жирным внутри буллета.
        if ":" in core and 0 < core.index(":") <= 80:
            head, tail = core.split(":", 1)
            _replace_runs(
                para,
                [
                    (f"{marker} ", False),
                    (head.strip() + ":", True),
                    (" " + tail.lstrip(), False),
                ],
            )
        else:
            _replace_runs(para, [(f"{marker} {core}", False)])
        eff_level = level if level > 0 else 1
        _apply_structure_style(para, indent_level=eff_level)
        return {"index": idx, "action": kind, "indent_level": eff_level, "old_text": old_text}

    if kind == "renumber":
        new_number = action.get("new_number", "")
        if not re.match(r"^\d+(?:\.\d+)*\.?$", new_number):
            return None
        if not new_number.endswith("."):
            new_number = new_number + "."
        mnum = NUMERIC_RE.match(old_text)
        if mnum:
            content = mnum.group(2).strip()
        else:
            content = old_text
        new_text = f"{new_number} {content}"
        _replace_runs(para, [(new_text, False)])
        _apply_structure_style(para, indent_level=level)
        return {
            "index": idx,
            "action": kind,
            "indent_level": level,
            "old_text": old_text,
            "new_text": new_text,
        }
    return None


def _safe_apply_actions(doc: Document, actions: list[dict[str, Any]]) -> dict[str, Any]:
    changed = 0
    applied: list[dict[str, Any]] = []
    seen_idx: set[int] = set()
    for raw in actions:
        norm = _normalize_action(raw)
        if not norm or norm["index"] in seen_idx:
            continue
        seen_idx.add(norm["index"])
        result = _apply_single_action(doc, norm)
        if result is not None:
            changed += 1
            applied.append(result)
    return {"changed_count": changed, "applied": applied}


# ============================================================
# ЭВРИСТИЧЕСКИЙ FALLBACK
# ============================================================

def _heuristic_structure_fix(doc: Document) -> dict[str, Any]:
    """
    Локальный fallback: если LLM молчит или недоступна — применим базовые
    эвристики: заголовки-секции жирним, подпункты под 'Преимущества:' делаем буллетами,
    и лечим "1 1 1 1" → "1 2 3 4" в локальных окнах.
    """
    changed = 0
    applied: list[dict[str, Any]] = []
    paras = doc.paragraphs
    texts = [(p.text or "").strip() for p in paras]

    nested_mode = False
    nested_expected = 1

    for i, para in enumerate(paras):
        txt = texts[i]
        if not txt:
            nested_mode = False
            nested_expected = 1
            continue
        m = NUMERIC_RE.match(txt)
        if not m:
            nested_mode = False
            nested_expected = 1
            continue
        num, content = m.group(1), m.group(2).strip()
        num_parts = num.split(".")
        is_top = len(num_parts) == 1
        try:
            num_last = int(num_parts[-1])
        except ValueError:
            num_last = -1

        prev = ""
        j = i - 1
        while j >= 0:
            if texts[j]:
                prev = texts[j]
                break
            j -= 1

        if prev and SUBLIST_PARENT_RE.search(prev) and is_top and num_last == 1:
            nested_mode = True
            nested_expected = 1

        if nested_mode and is_top and num_last == nested_expected:
            action = _normalize_action({"index": i, "action": "to_bullet", "indent_level": 2})
            if action:
                r = _apply_single_action(doc, action)
                if r:
                    changed += 1
                    applied.append(r)
            nested_expected += 1
            continue
        if nested_mode:
            nested_mode = False
            nested_expected = 1

        heading_like = (
            is_top
            and ":" not in content
            and len(content) <= 90
            and not content.endswith(".")
            and not re.search(
                r"\b(должен|должны|может|могут|нужно|нужны|является|это)\b",
                content.lower(),
            )
        )
        if heading_like:
            action = _normalize_action({"index": i, "action": "bold_full", "indent_level": 0})
            if action:
                r = _apply_single_action(doc, action)
                if r:
                    changed += 1
                    applied.append(r)
            continue

        if ":" in content:
            head = content.split(":", 1)[0].strip().lower()
            if len(head) <= 60 and any(
                kw in head for kw in (
                    "как работает", "вывод", "преимущест", "достоин",
                    "недостат", "плюс", "минус", "главная проблема",
                    "суть", "задача", "цель",
                )
            ):
                # «Суть:» / «Главная задача:» — обычный абзац (1.25 см), не блок 2.5 см как у буллетов.
                action = _normalize_action(
                    {"index": i, "action": "bold_prefix", "indent_level": 0}
                )
                if action:
                    r = _apply_single_action(doc, action)
                    if r:
                        changed += 1
                        applied.append(r)

    return {"changed_count": changed, "applied": applied}


# ============================================================
# ГЛАВНАЯ ФУНКЦИЯ
# ============================================================

def improve_doc_structure_with_ai(docx_path: str) -> dict[str, Any]:
    """
    Постобработка структуры уже готового DOCX через LLM.
    Возвращает краткий отчёт и при необходимости обновляет docx на месте.
    """
    t0 = time.time()
    enabled = str(getattr(settings, "AI_POSTPROCESS_ENABLED", "1")).lower() in {
        "1", "true", "yes", "on"
    }
    if not enabled:
        print("  [LLM] Постобработка структуры: отключена (AI_POSTPROCESS_ENABLED).")
        return {"enabled": False, "applied": False, "message": "AI postprocess disabled"}

    doc = Document(docx_path)
    _split_paragraphs_on_internal_newlines(doc)

    paragraphs = _collect_paragraphs(doc)
    document_head_preview = [
        {"global_index": i, "text_head": (paragraphs[i] or "")[:260]}
        for i in range(min(35, len(paragraphs)))
    ]
    window_size = int(getattr(settings, "AI_POSTPROCESS_WINDOW_SIZE", 50))
    overlap = int(getattr(settings, "AI_POSTPROCESS_WINDOW_OVERLAP", 5))
    max_windows = int(getattr(settings, "AI_POSTPROCESS_MAX_WINDOWS", 20))
    # Только «окна со списками» экономят токены, но тогда ввод без 1.2.3. не попадает в LLM,
    # если нумерация есть дальше по тексту — типичный конспект с таблицей/списком в конце.
    list_only = bool(getattr(settings, "AI_POSTPROCESS_LIST_ONLY_WINDOWS", False))
    if list_only:
        windows = _extract_windows(
            paragraphs, window_size=window_size, overlap=overlap, only_list_like=True
        )
        if not windows:
            windows = _extract_windows(
                paragraphs, window_size=window_size, overlap=overlap, only_list_like=False
            )
    else:
        windows = _extract_windows(
            paragraphs, window_size=window_size, overlap=overlap, only_list_like=False
        )
    if not windows:
        print("  [LLM] Постобработка структуры: нет абзацев для окна.")
        return {"enabled": True, "applied": False, "message": "No paragraphs"}
    windows_built_total = len(windows)
    windows_dropped_by_max = 0
    if len(windows) > max_windows:
        windows_dropped_by_max = len(windows) - max_windows
        windows = windows[:max_windows]

    extra = f", обрезано окон по лимиту: {windows_dropped_by_max}" if windows_dropped_by_max else ""
    print(
        f"  [LLM] Постобработка структуры: окон {len(windows)}, абзацев {len(paragraphs)}{extra}."
    )
    _print_llm_route_banner()

    run_ts = time.strftime("%Y%m%d_%H%M%S")
    digest_rows: list[tuple] = []
    windows_llm_trace: list[dict[str, Any]] = []

    t_llm_start = time.time()
    all_actions: list[dict[str, Any]] = []
    windows_meta: list[dict[str, Any]] = []
    errors: list[str] = []
    # Кто РЕАЛЬНО ответил (провайдер -> число окон с непустым ответом).
    providers_used: dict[str, int] = {}
    # Все провайдеры, которые хотя бы пробовались (для диагностики).
    providers_attempted: set[str] = set()

    inter_delay = float(getattr(settings, "AI_POSTPROCESS_INTER_WINDOW_SEC", 0) or 0)
    for wi, w in enumerate(windows):
        if wi > 0 and inter_delay > 0:
            time.sleep(inter_delay)
        prompt = _build_window_prompt(w)
        actions, err, meta = _call_llm(prompt)
        digest_rows.append((prompt, actions, err, meta))
        if _trace_enabled():
            windows_llm_trace.append(
                _build_window_llm_trace(wi, w, prompt, actions, err, meta)
            )
            if _trace_files_enabled():
                dbg_dir = getattr(settings, "AI_POSTPROCESS_DEBUG_DIR", "").strip()
                if not dbg_dir:
                    dbg_dir = os.path.join(str(settings.BASE_DIR), "debug_ai")
                inner_full = (meta or {}).get("llm_inner_text_full") or ""
                _write_window_trace_files(
                    dbg_dir, run_ts, wi, prompt, inner_full, err or ""
                )
        if _trace_console_enabled():
            first_gi = w.paragraphs[0][0] if w.paragraphs else -1
            print(
                f"  [LLM] trace окно {wi} абзацы [{w.start}–{w.end}] "
                f"первый_глоб._индекс={first_gi} промпт={len(prompt)} симв. "
                f"действий={len(actions)} err={'да' if err else 'нет'}"
            )
        prov = str(meta.get("provider", "") or "")
        if prov:
            providers_attempted.add(prov)
            if actions:
                providers_used[prov] = providers_used.get(prov, 0) + 1
        meta_for_log = dict(meta)
        meta_for_log.pop("llm_inner_text_full", None)
        att = meta_for_log.get("attempts")
        if isinstance(att, list):
            meta_for_log["attempts"] = [
                {k: v for k, v in (a or {}).items() if k != "llm_inner_text_full"}
                if isinstance(a, dict)
                else a
                for a in att
            ]
        windows_meta.append({
            "range": [w.start, w.end],
            "paragraphs_in_window": len(w.paragraphs),
            "actions_count": len(actions),
            "err": err,
            "meta": meta_for_log,
        })
        if err:
            errors.append(err)
        if actions:
            all_actions.extend(actions)

    llm_ms = int((time.time() - t_llm_start) * 1000)

    trace_extra: dict[str, Any] = {
        "run_ts": run_ts,
        "document_head_preview": document_head_preview,
        "paragraphs_collect_note": (
            "Индексы — только верхний поток body (python-docx: doc.paragraphs). "
            "Абзацы внутри таблииц Word сюда не попадают; их LLM не видит и не размечает."
        ),
        "windows_llm_digest": _windows_digest(windows, digest_rows),
    }
    if _trace_enabled() and windows_llm_trace:
        trace_extra["windows_llm_trace"] = windows_llm_trace

    windows_diag = {
        "windows_built_total": windows_built_total,
        "windows_sent_to_llm": len(windows),
        "windows_dropped_by_max": windows_dropped_by_max,
        "max_windows_setting": max_windows,
        "windows_with_errors": [
            {"range": wm["range"], "err": (wm.get("err") or "")[:1200]}
            for wm in windows_meta if wm.get("err")
        ],
    }

    def _actual_source() -> str:
        """Провайдер, который реально дал хоть один непустой ответ.
        Если таких несколько (fallback сработал в разных окнах) —
        возвращаем их через '+', например 'gemini+ollama'."""
        if not providers_used:
            return ""
        ordered = sorted(providers_used.items(), key=lambda kv: (-kv[1], kv[0]))
        return "+".join(name for name, _ in ordered)

    # Дедуп по индексу — последнее действие для индекса побеждает.
    dedup: dict[int, dict[str, Any]] = {}
    for a in all_actions:
        if not isinstance(a, dict):
            continue
        try:
            k = int(a.get("index"))
        except Exception:
            continue
        dedup[k] = a
    actions_unique = list(dedup.values())

    if not actions_unique:
        use_heuristic = str(
            getattr(settings, "AI_POSTPROCESS_HEURISTIC_FALLBACK", "1")
        ).lower() in {"1", "true", "yes", "on"}
        if use_heuristic:
            heuristic = _heuristic_structure_fix(doc)
            if heuristic["changed_count"] > 0:
                doc.save(docx_path)
                total_ms = int((time.time() - t0) * 1000)
                _write_debug_log({
                    "docx_path": docx_path,
                    "status": "heuristic_applied",
                    "provider_primary": _resolve_provider(),
                    "provider_actual": "heuristic",
                    "providers_used": providers_used,
                    "providers_attempted": sorted(providers_attempted),
                    "timing_ms": {"total": total_ms, "llm": llm_ms},
                    "paragraph_count": len(paragraphs),
                    "window_count": len(windows),
                    "windows_diagnosis": windows_diag,
                    "changed_count": heuristic["changed_count"],
                    "changes_preview": heuristic["applied"][:30],
                    "windows_debug": windows_meta,
                    "errors": errors,
                    **trace_extra,
                }, run_ts=run_ts)
                print(
                    f"  [LLM] Нет ответа с правками; применена эвристика: "
                    f"{heuristic['changed_count']} измен."
                )
                if errors:
                    print(
                        "  [LLM] Запросы к модели не удались (подробности в debug_ai/*.json при AI_POSTPROCESS_DEBUG). "
                        "Эвристика только жирнит заголовки — буллеты и списки, как у LLM, не восстанавливает. "
                        "Повторите при стабильном интернете или GEMINI_PROXY / OpenRouter."
                    )
                return {
                    "enabled": True,
                    "applied": True,
                    "source": "heuristic",
                    "providers_attempted": sorted(providers_attempted),
                    "changed_count": heuristic["changed_count"],
                    "changes": heuristic["applied"][:40],
                    "timing_ms": {"total": total_ms, "llm": llm_ms},
                }

        msg = "No suggested changes from LLM"
        if errors:
            msg = f"LLM errors: {'; '.join(sorted(set(errors)))[:500]}"
        total_ms = int((time.time() - t0) * 1000)
        print(f"  [LLM] Без правок. {msg[:400]}")
        _write_debug_log({
            "docx_path": docx_path,
            "status": "no_actions",
            "provider_primary": _resolve_provider(),
            "providers_attempted": sorted(providers_attempted),
            "message": msg,
            "timing_ms": {"total": total_ms, "llm": llm_ms},
            "paragraph_count": len(paragraphs),
            "window_count": len(windows),
            "windows_diagnosis": windows_diag,
            "windows_debug": windows_meta,
            **trace_extra,
        }, run_ts=run_ts)
        return {
            "enabled": True,
            "applied": False,
            "source": "",
            "providers_attempted": sorted(providers_attempted),
            "message": msg,
        }

    result = _safe_apply_actions(doc, actions_unique)
    total_ms = int((time.time() - t0) * 1000)
    source = _actual_source() or _resolve_provider()
    _write_debug_log({
        "docx_path": docx_path,
        "status": "ok",
        "provider_primary": _resolve_provider(),
        "provider_actual": source,
        "providers_used": providers_used,
        "providers_attempted": sorted(providers_attempted),
        "timing_ms": {"total": total_ms, "llm": llm_ms},
        "paragraph_count": len(paragraphs),
        "window_count": len(windows),
        "windows_diagnosis": windows_diag,
        "action_count": len(actions_unique),
        "changed_count": result["changed_count"],
        "changes_preview": result["applied"][:30],
        "windows_debug": windows_meta,
        **trace_extra,
    }, run_ts=run_ts)
    if result["changed_count"] > 0:
        doc.save(docx_path)
        print(f"  [LLM] Применено изменений {result['changed_count']} (источник: {source}).")
        return {
            "enabled": True,
            "applied": True,
            "source": source,
            "providers_used": providers_used,
            "changed_count": result["changed_count"],
            "changes": result["applied"][:40],
            "timing_ms": {"total": total_ms, "llm": llm_ms},
        }
    print(
        f"  [LLM] Предложения есть, безопасно применить не удалось (источник: {source})."
    )
    return {
        "enabled": True,
        "applied": False,
        "source": source,
        "providers_used": providers_used,
        "message": "No safe changes",
        "timing_ms": {"total": total_ms, "llm": llm_ms},
    }
