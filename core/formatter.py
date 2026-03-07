"""
Форматтер документов по ГОСТ.
Создаёт новый .docx с правильным оформлением.
Поддержка рамок: VML-фигуры в body → перенос в header для повтора на всех страницах.
"""

import os
import re

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree


GOST = {
    "margin_left": Cm(3),
    "margin_right": Cm(1.5),
    "margin_top": Cm(2),
    "margin_bottom": Cm(2),
    "font": "Times New Roman",
    "font_size": Pt(14),
    "font_size_caption": Pt(12),
    "line_spacing": 1.5,
    "indent": Cm(1.25),
}


def set_font(run, bold=False, size=None):
    """Устанавливает шрифт для run"""
    run.font.name = GOST["font"]
    run.font.size = size or GOST["font_size"]
    run.font.bold = bold
    try:
        run._element.rPr.rFonts.set(qn('w:eastAsia'), GOST["font"])
    except:
        pass


# ================================================================
# РАБОТА С РАМКОЙ
# ================================================================

def _has_graphic_content(xml_str):
    """
    Проверяет, содержит ли XML графические объекты (VML, Drawing).
    Это рамка — линии, прямоугольники, текстовые блоки штампа.
    """
    markers = [
        '<w:drawing',
        '<w:pict',
        '<mc:AlternateContent',
        '<v:shape', '<v:rect', '<v:line', '<v:group',
        '<v:roundrect', '<v:oval', '<v:textbox',
        '<wps:wsp', '<wpg:wgp',
        '<wp:anchor', '<wp:inline',
    ]
    return any(m in xml_str for m in markers)


def _header_already_has_frame(section):
    """Проверяет, есть ли уже рамка в header (для шаблонов с header-подходом)"""
    try:
        header = section.header
        if header.is_linked_to_previous:
            return False
        h_xml = etree.tostring(header._element).decode()
        return _has_graphic_content(h_xml)
    except:
        return False


def prepare_template(template_path):
    """
    Подготовка шаблона рамки.

    Рамка может быть реализована двумя способами:
    1. Уже в header/footer (классический) → просто очищаем body
    2. VML-фигуры в body параграфах (ваш случай!) → переносим в header

    Перенос в header делает рамку видимой на КАЖДОЙ странице.
    """
    print(f"\n  {'─'*50}")
    print(f"  🖼️ ПОДГОТОВКА РАМКИ")
    print(f"  {'─'*50}")
    print(f"  📂 Файл: {os.path.basename(template_path)}")
    print(f"  📂 Размер: {os.path.getsize(template_path)} байт")

    doc = Document(template_path)
    body = doc.element.body
    section = doc.sections[0]

    # ─── СЛУЧАЙ 1: Рамка уже в header ───
    if _header_already_has_frame(section):
        print(f"  ✅ Рамка уже в header — очищаем body")
        for child in list(body):
            if child.tag in (qn('w:p'), qn('w:tbl')):
                body.remove(child)
        return doc

    # ─── СЛУЧАЙ 2: Рамка в body (VML в параграфах) ───
    print(f"  📋 Рамка в body — переносим в header...")

    # Анализируем каждый параграф
    total_paras = 0
    vml_runs = []  # Все VML-элементы для переноса в header

    for child in list(body):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

        if tag == 'sectPr':
            continue  # sectPr не трогаем

        if tag == 'p':
            total_paras += 1
            p_xml = etree.tostring(child).decode()
            has_vml = _has_graphic_content(p_xml)

            if has_vml:
                # Извлекаем все не-pPr дочерние элементы (runs с VML)
                for sub in list(child):
                    if sub.tag != qn('w:pPr'):
                        child.remove(sub)
                        vml_runs.append(sub)
                texts = ''.join(child.itertext()).strip()[:40]
                print(f"     ✅ VML-параграф → {len(vml_runs)} элементов (text='{texts}')")

            body.remove(child)

        elif tag == 'tbl':
            body.remove(child)

    print(f"  📊 Параграфов в шаблоне: {total_paras}")
    print(f"  📊 VML-элементов извлечено: {len(vml_runs)}")

    if not vml_runs:
        print(f"  ⚠️ VML-элементы не найдены!")
        print(f"  ⚠️ Используем только настройки полей из шаблона")
        return doc

    # ─── Создаём header с рамкой ───
    header = section.header
    header.is_linked_to_previous = False

    h_elem = header._element
    # Удаляем дефолтный пустой параграф
    for child in list(h_elem):
        h_elem.remove(child)

    # Один параграф с МИНИМАЛЬНОЙ высотой, содержащий ВСЕ VML-фигуры
    # Высота 1 twip ≈ 0.018мм — невидима, но VML отрисуется
    # на своих абсолютных позициях на странице
    h_para = OxmlElement('w:p')

    h_pPr = OxmlElement('w:pPr')
    h_spacing = OxmlElement('w:spacing')
    h_spacing.set(qn('w:line'), '1')        # 1 twip — минимум
    h_spacing.set(qn('w:lineRule'), 'exact')
    h_spacing.set(qn('w:before'), '0')
    h_spacing.set(qn('w:after'), '0')
    h_pPr.append(h_spacing)

    # Размер шрифта 1pt на случай если есть невидимый текст
    h_rPr = OxmlElement('w:rPr')
    h_sz = OxmlElement('w:sz')
    h_sz.set(qn('w:val'), '2')  # 1pt = 2 half-points
    h_rPr.append(h_sz)
    h_szCs = OxmlElement('w:szCs')
    h_szCs.set(qn('w:val'), '2')
    h_rPr.append(h_szCs)
    h_pPr.append(h_rPr)

    h_para.append(h_pPr)

    # Добавляем все VML-элементы в один параграф
    for run in vml_runs:
        h_para.append(run)

    h_elem.append(h_para)

    # Проверка
    verify_xml = etree.tostring(h_elem).decode()
    has_vml_in_header = _has_graphic_content(verify_xml)
    print(f"  ✅ Header создан: VML={has_vml_in_header}, размер XML={len(verify_xml)} байт")

    # Проверяем что sectPr получил headerReference
    try:
        sect_xml = etree.tostring(section._sectPr).decode()
        has_href = 'headerReference' in sect_xml
        print(f"  📋 sectPr headerReference: {has_href}")
    except:
        pass

    return doc


# ================================================================
# ВСПОМОГАТЕЛЬНЫЕ
# ================================================================

def _move_sectpr_to_end(doc):
    """Перемещает sectPr в конец body (требование OOXML)"""
    body = doc.element.body
    moved = 0
    for child in list(body):
        if child.tag == qn('w:sectPr') or child.tag.endswith('}sectPr'):
            body.remove(child)
            body.append(child)
            moved += 1
    if moved:
        print(f"  🔧 sectPr → конец body ({moved} шт.)")


def _set_margins(doc):
    """Устанавливает поля страницы по ГОСТ (без рамки)"""
    sec = doc.sections[0]
    sec.left_margin = GOST["margin_left"]
    sec.right_margin = GOST["margin_right"]
    sec.top_margin = GOST["margin_top"]
    sec.bottom_margin = GOST["margin_bottom"]


# ================================================================
# ГЛАВНАЯ ФУНКЦИЯ
# ================================================================

def create_gost_document(elements, output_path, template_path=None):
    """
    Создаёт новый документ по ГОСТ из классифицированных элементов.

    Если указана рамка (template_path):
    - VML-фигуры рамки переносятся в header (повтор на каждой странице)
    - Поля страницы берутся из шаблона рамки (не ГОСТ)
    """

    print(f"\n{'='*60}")
    print(f"🔨 СОЗДАНИЕ ДОКУМЕНТА ПО ГОСТ")
    print(f"{'='*60}")

    use_template = False

    if template_path and os.path.exists(template_path):
        print(f"  🖼️ Рамка: {os.path.basename(template_path)}")
        try:
            doc = prepare_template(template_path)
            use_template = True
            # НЕ меняем поля — шаблон имеет свои, подогнанные под рамку
            _set_margins(doc)
            print(f"  ✅ Рамка применена!")
        except Exception as e:
            print(f"  ❌ Ошибка рамки: {e}")
            import traceback
            traceback.print_exc()
            doc = Document()
            _set_margins(doc)
    else:
        if template_path:
            print(f"  ❌ Файл рамки не найден: {template_path}")
        doc = Document()
        _set_margins(doc)

    # ─── Счётчики заголовков ───
    chapter_num = 0
    section_num = 0
    subsection_num = 0
    elements_written = 0

    # ─── Добавляем элементы ───
    for elem in elements:
        text = elem.get("text", "")
        etype = elem["type"]

        if not text and etype != "image":
            continue

        # === IMAGE ===
        if etype == "image":
            for img_path in elem.get("image_paths", []):
                if os.path.exists(img_path):
                    para = doc.add_paragraph()
                    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    para.paragraph_format.line_spacing = GOST["line_spacing"]
                    para.paragraph_format.first_line_indent = Cm(0)
                    para.paragraph_format.space_before = Pt(6)
                    para.paragraph_format.space_after = Pt(6)
                    run = para.add_run()
                    try:
                        run.add_picture(img_path, width=Cm(16))
                        elements_written += 1
                    except Exception as e:
                        print(f"  ⚠️ Ошибка изображения: {e}")
            continue

        para = doc.add_paragraph()
        para.paragraph_format.line_spacing = GOST["line_spacing"]
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.space_before = Pt(0)

        # === FIGURE_CAPTION ===
        if etype == "figure_caption":
            clean = re.sub(r'^рисунок\s+\d+\s*[–\-—.:]\s*', '', text, flags=re.IGNORECASE)
            fig_num = elem.get("figure_num", 1)
            run = para.add_run(f"Рисунок {fig_num} – {clean}")
            set_font(run, bold=False, size=GOST["font_size_caption"])
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.first_line_indent = Cm(0)
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(12)

        # === HEADING_CAPS ===
        elif etype == "heading_caps":
            run = para.add_run(text.upper())
            set_font(run, bold=True)
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(24)
            para.paragraph_format.space_after = Pt(12)
            para.paragraph_format.first_line_indent = Cm(0)

        # === HEADING_TITLE ===
        elif etype == "heading_title":
            run = para.add_run(text)
            set_font(run, bold=True)
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(18)
            para.paragraph_format.space_after = Pt(12)
            para.paragraph_format.first_line_indent = Cm(0)

        # === HEADING_CHAPTER ===
        elif etype == "heading_chapter":
            chapter_num += 1
            section_num = 0
            subsection_num = 0
            if not re.match(r'^\d+\.', text):
                text = f"{chapter_num}. {text}"
            run = para.add_run(text)
            set_font(run, bold=True)
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(6)
            para.paragraph_format.first_line_indent = GOST["indent"]

        # === HEADING_SECTION ===
        elif etype == "heading_section":
            if chapter_num == 0:
                chapter_num = 1
            section_num += 1
            subsection_num = 0
            if not re.match(r'^\d+\.\d+', text):
                text = f"{chapter_num}.{section_num}. {text}"
            run = para.add_run(text)
            set_font(run, bold=True)
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(6)
            para.paragraph_format.first_line_indent = GOST["indent"]

        # === HEADING_SUBSECTION ===
        elif etype == "heading_subsection":
            if chapter_num == 0:
                chapter_num = 1
            if section_num == 0:
                section_num = 1
            subsection_num += 1
            if not re.match(r'^\d+\.\d+\.\d+', text):
                text = f"{chapter_num}.{section_num}.{subsection_num}. {text}"
            run = para.add_run(text)
            set_font(run, bold=True)
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(6)
            para.paragraph_format.first_line_indent = GOST["indent"]

        # === NUMBERED_ITEM ===
        elif etype == "numbered_item":
            num = elem.get("number", "")
            full_text = f"{num}. {text}" if num else text
            run = para.add_run(full_text)
            set_font(run, bold=False)
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.first_line_indent = GOST["indent"]

        # === LIST_ITEM ===
        elif etype == "list_item":
            clean = re.sub(r'^[•\-–—\*●○◦]\s*', '', text)
            run = para.add_run("— " + clean)
            set_font(run, bold=False)
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.first_line_indent = GOST["indent"]

        # === PARAGRAPH ===
        else:
            run = para.add_run(text)
            set_font(run, bold=False)
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.first_line_indent = GOST["indent"]

        elements_written += 1

    # ─── Финализация ───
    _move_sectpr_to_end(doc)

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)

    print(f"\n  ✅ Записано: {elements_written} элементов")
    print(f"  📂 Файл: {output_path} ({os.path.getsize(output_path)} байт)")

    # Верификация: проверяем что header есть в результате
    if use_template:
        import zipfile
        with zipfile.ZipFile(output_path, 'r') as z:
            headers = [f for f in z.namelist() if 'header' in f.lower()]
            print(f"  🔍 Headers в результате: {headers}")
            if headers:
                print(f"  ✅ Рамка должна быть на каждой странице!")
            else:
                print(f"  ❌ Header не создан — рамка может не работать")

    return output_path