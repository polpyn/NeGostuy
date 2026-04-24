"""
API views для проекта НеГостуй

Эндпоинты:
- POST /api/upload/          — загрузка документа
- GET  /api/status/<id>/     — статус обработки
- GET  /api/report/<id>/     — отчёт об ошибках
- GET  /api/download/<id>/   — скачивание результата
- POST /api/auth/register/   — регистрация
- POST /api/auth/login/      — вход
- GET  /api/health/          — проверка сервера
"""

import os
import time
import json
import tempfile
from copy import deepcopy
from urllib.parse import quote

from django.conf import settings
from django.http import FileResponse, JsonResponse, HttpResponse
from django.views.decorators.csrf import csrf_exempt
from django.contrib.auth.models import User
from django.contrib.auth import authenticate, login, logout

from docx import Document as DocxDocument
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION_START
from docxcompose.composer import Composer
from rest_framework.decorators import api_view, permission_classes
from rest_framework.permissions import AllowAny, IsAuthenticated
from rest_framework.response import Response
from rest_framework import status
try:
    from celery.result import AsyncResult
except Exception:
    AsyncResult = None

from .models import Document, ProcessingResult, GOSTTemplate, ErrorStatistic


def _attachment_content_disposition(filename: str) -> str:
    """
    Content-Disposition для скачивания: ASCII filename + RFC 5987 для кириллицы и прочего Unicode.
    Иначе заголовок не кодируется в latin-1 или клиенты ломают имя файла.
    """
    filename = (filename or "download.docx").replace('"', "_")
    try:
        filename.encode("latin-1")
        return f'attachment; filename="{filename}"'
    except UnicodeEncodeError:
        ascii_fallback = (
            filename.encode("ascii", "ignore").decode("ascii").strip(". ") or "download.docx"
        )
        return (
            f'attachment; filename="{ascii_fallback}"; '
            f"filename*=UTF-8''{quote(filename, safe='')}"
        )
from core.analyzer import analyze_document
try:
    from .tasks import process_document_task
except Exception:
    process_document_task = None


def _value_from_query_mapping(mapping, key: str):
    """Одно непустое значение по ключу (QueryDict или совместимый объект)."""
    if hasattr(mapping, "getlist"):
        for v in mapping.getlist(key):
            if v is not None and str(v).strip():
                return str(v).strip()
    if hasattr(mapping, "get"):
        v = mapping.get(key)
        if v is not None and str(v).strip():
            return str(v).strip()
    return None


def _multipart_plain_value(request, *keys: str) -> str:
    """
    Первое непустое текстовое значение из multipart по списку ключей.

    У DRF тело multipart читается парсером в request.data / request.POST (обёртка).
    Внутренний request._request.POST после этого часто пустой — нельзя опираться на него первым.
    """
    seen: set[int] = set()
    sources: list = []

    drf_data = getattr(request, "data", None)
    if drf_data is not None:
        sources.append(drf_data)

    wrapped_post = getattr(request, "POST", None)
    if wrapped_post is not None:
        sources.append(wrapped_post)

    if hasattr(request, "_request"):
        inner = request._request.POST
        sources.append(inner)

    for mapping in sources:
        if mapping is None:
            continue
        mid = id(mapping)
        if mid in seen:
            continue
        seen.add(mid)
        for key in keys:
            got = _value_from_query_mapping(mapping, key)
            if got:
                return got
    return ""


# ============================================================
# HEALTH CHECK
# ============================================================

@api_view(['GET'])
@permission_classes([AllowAny])
def health_check(request):
    """Проверка работоспособности сервера"""
    return Response({
        'status': 'ok',
        'message': 'НеГостуй API работает',
        'version': '1.0'
    })


# ============================================================
# ЗАГРУЗКА И ОБРАБОТКА ДОКУМЕНТА
# ============================================================

@api_view(['POST'])
@permission_classes([AllowAny])
def upload_document(request):
    """
    POST /api/upload/
    Принимает .docx файл + необязательную рамку (template).
    """

    if 'file' not in request.FILES:
        return Response(
            {'error': 'Файл не найден. Отправьте файл в поле "file"'},
            status=status.HTTP_400_BAD_REQUEST
        )

    file = request.FILES['file']

    # Валидация документа
    if not file.name.endswith('.docx'):
        return Response({'error': 'Разрешены только .docx'}, status=status.HTTP_400_BAD_REQUEST)
    if file.size > 16 * 1024 * 1024:
        return Response({'error': 'Максимум 16 МБ'}, status=status.HTTP_400_BAD_REQUEST)
    if file.size == 0:
        return Response({'error': 'Файл пустой'}, status=status.HTTP_400_BAD_REQUEST)

    # Рамка: загрузка или ramka.docx в корне проекта
    template_file = request.FILES.get('template', None)
    template_path = None
    template_is_temp = False

    if template_file:
        if not template_file.name.endswith('.docx'):
            return Response({'error': 'Рамка должна быть .docx'}, status=status.HTTP_400_BAD_REQUEST)

        # Сохраняем рамку во временный файл
        import tempfile
        tmp_template = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        for chunk in template_file.chunks():
            tmp_template.write(chunk)
        tmp_template.close()
        template_path = tmp_template.name
        template_is_temp = True
        print(f"Рамка сохранена: {template_path}")
    else:
        default_frame = os.path.join(settings.BASE_DIR, 'ramka.docx')
        if os.path.isfile(default_frame):
            template_path = default_frame
            print(f"Рамка по умолчанию: {template_path}")

    zachet_number = _multipart_plain_value(
        request, "student_id", "zachet_number", "frame_zachet"
    )
    if zachet_number:
        print(f"Номер зачётной книжки (принят): {zachet_number!r}")
    else:
        print("Номер зачётной книжки: не передан или пустой (проверьте поля student_id / zachet_number)")

    work_type = request.data.get('work_type', 'coursework')

    # Сохраняем документ в БД и ставим в очередь Celery
    doc = Document.objects.create(
        user=request.user if request.user.is_authenticated else None,
        filename=file.name,
        original_file=file,
        file_size=file.size,
        work_type=work_type,
        status='queued',
        progress=5,
    )

    try:
        if settings.ENABLE_ASYNC_PROCESSING and process_document_task is not None:
            try:
                async_result = process_document_task.delay(
                    doc.id,
                    template_path,
                    template_is_temp,
                    zachet_number or "",
                )
                doc.task_id = async_result.id
                doc.save(update_fields=['task_id'])
                return Response({
                    'success': True,
                    'document_id': doc.id,
                    'document_name': doc.filename,
                    'status': 'queued',
                    'task_id': async_result.id,
                    'summary': {
                        'has_template': bool(template_path),
                        'zachet_received': zachet_number,
                    },
                    'elements': [],
                })
            except Exception as celery_exc:
                # Локальный fallback: Redis/Celery недоступен — обрабатываем синхронно.
                print(f"Celery недоступен, fallback на синхронную обработку: {celery_exc}")
        elif not settings.ENABLE_ASYNC_PROCESSING:
            print("Асинхронная обработка отключена (ENABLE_ASYNC_PROCESSING=0), запускаем синхронно.")

        # Fallback: если Celery не установлен/не запущен, обрабатываем синхронно
        doc.status = 'parsing'
        doc.progress = 10
        doc.save(update_fields=['status', 'progress'])
        start_time = time.time()
        result_data = analyze_document(
            doc.original_file.path,
            template_path=template_path,
            zachet_number=zachet_number or None,
        )
        processing_time = time.time() - start_time
        proc_result = ProcessingResult.objects.create(
            document=doc,
            report_json=result_data.get('report', {}),
            total_elements=result_data.get('total_elements', 0),
            errors_count=result_data.get('errors_count', 0),
            warnings_count=result_data.get('warnings_count', 0),
            headings_count=result_data.get('headings_count', 0),
            paragraphs_count=result_data.get('paragraphs_count', 0),
            lists_count=result_data.get('lists_count', 0),
            images_count=result_data.get('images_count', 0),
            grade=result_data.get('grade', ''),
            processing_time=processing_time,
        )
        if result_data.get('output_path') and os.path.exists(result_data['output_path']):
            from django.core.files import File
            with open(result_data['output_path'], 'rb') as f:
                proc_result.output_file.save(f"GOST_{doc.filename}", File(f))
        doc.status = 'completed'
        doc.progress = 100
        doc.save(update_fields=['status', 'progress'])
        return Response({
            'success': True,
            'document_id': doc.id,
            'document_name': doc.filename,
            'status': 'completed',
            'summary': {
                'total_elements': proc_result.total_elements,
                'headings': proc_result.headings_count,
                'lists': proc_result.lists_count,
                'texts': proc_result.paragraphs_count,
                'images': proc_result.images_count,
                'errors_count': proc_result.errors_count,
                'warnings_count': proc_result.warnings_count,
                'perfect_elements': proc_result.total_elements - proc_result.errors_count,
                'elements_with_errors': proc_result.errors_count,
                'grade': proc_result.grade,
                'processing_time': round(processing_time, 2),
                'has_template': bool(template_path),
                'zachet_received': zachet_number,
                'ai_postprocess': result_data.get('ai_postprocess', {}),
            },
            'elements': result_data.get('elements_detail', []),
        })
    except Exception as e:
        doc.status = 'error'
        doc.progress = 100
        doc.save(update_fields=['status', 'progress'])
        if template_is_temp and template_path and os.path.exists(template_path):
            try:
                os.unlink(template_path)
            except OSError:
                pass
        return Response(
            {'success': False, 'error': str(e)},
            status=status.HTTP_500_INTERNAL_SERVER_ERROR
        )

# ============================================================
# СТАТУС ОБРАБОТКИ
# ============================================================

@api_view(['GET'])
@permission_classes([AllowAny])
def document_status(request, doc_id):
    """
    GET /api/status/<id>/
    Возвращает текущий статус обработки документа.
    """
    try:
        doc = Document.objects.get(id=doc_id)
        task_state = None
        if doc.task_id and AsyncResult is not None:
            try:
                task_state = AsyncResult(doc.task_id).state
            except Exception:
                task_state = None
        return Response({
            'document_id': doc.id,
            'filename': doc.filename,
            'status': doc.status,
            'progress': doc.progress,
            'status_display': doc.get_status_display(),
            'task_id': doc.task_id,
            'task_state': task_state,
        })
    except Document.DoesNotExist:
        return Response(
            {'error': 'Документ не найден'},
            status=status.HTTP_404_NOT_FOUND
        )


# ============================================================
# ОТЧЁТ ОБ ОШИБКАХ
# ============================================================

@api_view(['GET'])
@permission_classes([AllowAny])
def document_report(request, doc_id):
    """
    GET /api/report/<id>/
    Возвращает подробный отчёт об ошибках.
    """
    try:
        doc = Document.objects.get(id=doc_id)

        if doc.status != 'completed':
            return Response({
                'error': 'Документ ещё обрабатывается',
                'status': doc.status,
                'progress': doc.progress,
            })

        result = doc.result

        return Response({
            'document_id': doc.id,
            'filename': doc.filename,
            'report': result.report_json,
            'summary': {
                'total_elements': result.total_elements,
                'errors_count': result.errors_count,
                'warnings_count': result.warnings_count,
                'grade': result.grade,
                'processing_time': result.processing_time,
                'ai_postprocess': (result.report_json or {}).get('ai_postprocess', {}),
            }
        })

    except Document.DoesNotExist:
        return Response(
            {'error': 'Документ не найден'},
            status=status.HTTP_404_NOT_FOUND
        )
    except ProcessingResult.DoesNotExist:
        return Response(
            {'error': 'Результат ещё не готов'},
            status=status.HTTP_404_NOT_FOUND
        )


# ============================================================
# СКАЧИВАНИЕ РЕЗУЛЬТАТА
# ============================================================

@api_view(['GET'])
@permission_classes([AllowAny])
def download_result(request, doc_id):
    """
    GET /api/download/<id>/
    Возвращает исправленный .docx файл.
    """
    try:
        doc = Document.objects.get(id=doc_id)

        if doc.status != 'completed':
            return Response(
                {'error': 'Документ ещё обрабатывается'},
                status=status.HTTP_400_BAD_REQUEST
            )

        result = doc.result

        if not result.output_file:
            return Response(
                {'error': 'Исправленный файл не найден'},
                status=status.HTTP_404_NOT_FOUND
            )

        return FileResponse(
            result.output_file.open('rb'),
            as_attachment=True,
            filename=f"GOST_{doc.filename}"
        )

    except (Document.DoesNotExist, ProcessingResult.DoesNotExist):
        return Response(
            {'error': 'Документ не найден'},
            status=status.HTTP_404_NOT_FOUND
        )


# ============================================================
# АУТЕНТИФИКАЦИЯ
# ============================================================

@api_view(['POST'])
@permission_classes([AllowAny])
def register_user(request):
    """
    POST /api/auth/register/
    Регистрация нового пользователя.
    """
    username = request.data.get('username', '').strip()
    email = request.data.get('email', '').strip()
    password = request.data.get('password', '')

    # Валидация
    if not username:
        return Response(
            {'error': 'Укажите имя пользователя'},
            status=status.HTTP_400_BAD_REQUEST
        )
    if not password or len(password) < 6:
        return Response(
            {'error': 'Пароль должен содержать минимум 6 символов'},
            status=status.HTTP_400_BAD_REQUEST
        )
    if User.objects.filter(username=username).exists():
        return Response(
            {'error': 'Пользователь с таким именем уже существует'},
            status=status.HTTP_400_BAD_REQUEST
        )
    if email and User.objects.filter(email=email).exists():
        return Response(
            {'error': 'Email уже используется'},
            status=status.HTTP_400_BAD_REQUEST
        )

    user = User.objects.create_user(
        username=username,
        email=email,
        password=password
    )

    return Response({
        'success': True,
        'message': 'Регистрация успешна',
        'user_id': user.id,
        'username': user.username,
    }, status=status.HTTP_201_CREATED)


@api_view(['POST'])
@permission_classes([AllowAny])
def login_user(request):
    """
    POST /api/auth/login/
    Вход пользователя.
    """
    username = request.data.get('username', '')
    password = request.data.get('password', '')

    user = authenticate(request, username=username, password=password)

    if user is not None:
        login(request, user)
        return Response({
            'success': True,
            'message': 'Вход выполнен',
            'user_id': user.id,
            'username': user.username,
        })
    else:
        return Response(
            {'error': 'Неверный логин или пароль'},
            status=status.HTTP_401_UNAUTHORIZED
        )


@api_view(['POST'])
def logout_user(request):
    """POST /api/auth/logout/"""
    logout(request)
    return Response({'success': True, 'message': 'Выход выполнен'})


# ============================================================
# ВСПОМОГАТЕЛЬНЫЕ
# ============================================================

# ============================================================
# ТИТУЛЬНЫЙ ЛИСТ
# ============================================================

def _merge_title_and_report(title_path: str, report_path: str, output_path: str) -> None:
    """
    Склеивает документы: [титульник] + [отчёт].
    Используем штатный merger (docxcompose), чтобы корректно сохранить секции:
    - рамка/колонтитулы титульника остаются на первой части;
    - рамка отчёта применяется к страницам отчёта;
    - не ломается геометрия первой страницы (год не "уезжает").
    """
    # Сборка через Composer, чтобы сохранить relationship'ы (картинки, стили и т.д.).
    # Перед append гарантируем отдельную секцию после титульника.
    title_doc = DocxDocument(title_path)
    report_doc = DocxDocument(report_path)
    title_doc.add_section(WD_SECTION_START.NEW_PAGE)
    composer = Composer(title_doc)
    composer.append(report_doc)
    composer.save(output_path)

    # Пост-фикс секций/колонтитулов и нумерации.
    merged_doc = DocxDocument(output_path)
    if len(merged_doc.sections) < 2:
        merged_doc.add_section(WD_SECTION_START.NEW_PAGE)

    report_src = report_doc.sections[0]

    def _copy_hdrftr(src_hdrftr, dst_hdrftr):
        dst_el = dst_hdrftr._element
        for child in list(dst_el):
            dst_el.remove(child)
        for child in list(src_hdrftr._element):
            dst_el.append(deepcopy(child))

    def _strip_page_fields(hdrftr):
        # Убираем только PAGE-поля в fldSimple (именно так вставляется {{list}}),
        # чтобы на титульнике не печаталась "1".
        el = hdrftr._element
        for fld in list(el.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldSimple')):
            instr = fld.get(qn('w:instr')) or ''
            if 'PAGE' in instr:
                parent = fld.getparent()
                if parent is not None:
                    parent.remove(fld)

    def _set_section_page_start(sec, start_value: int):
        sect_pr = sec._sectPr
        pg_num = sect_pr.find(qn('w:pgNumType'))
        if pg_num is None:
            pg_num = OxmlElement('w:pgNumType')
            sect_pr.append(pg_num)
        pg_num.set(qn('w:start'), str(start_value))

    # Все секции отчётной части должны иметь рамку отчёта и не наследовать титульник.
    for sec in merged_doc.sections[1:]:
        sec.different_first_page_header_footer = report_src.different_first_page_header_footer
        sec.header.is_linked_to_previous = False
        sec.footer.is_linked_to_previous = False
        sec.even_page_header.is_linked_to_previous = False
        sec.even_page_footer.is_linked_to_previous = False
        sec.first_page_header.is_linked_to_previous = False
        sec.first_page_footer.is_linked_to_previous = False
        _copy_hdrftr(report_src.header, sec.header)
        _copy_hdrftr(report_src.footer, sec.footer)
        _copy_hdrftr(report_src.even_page_header, sec.even_page_header)
        _copy_hdrftr(report_src.even_page_footer, sec.even_page_footer)
        _copy_hdrftr(report_src.first_page_header, sec.first_page_header)
        _copy_hdrftr(report_src.first_page_footer, sec.first_page_footer)

    # Нумерация должна стартовать с 2 во всём документе.
    # На титульнике номер удалён, поэтому визуально это влияет только на отчёт.
    for sec in merged_doc.sections:
        _set_section_page_start(sec, 2)

    # На титульнике убираем номер страницы, чтобы последовательность
    # в отчётной части начиналась с "2" без лишней первой "1".
    title_first = merged_doc.sections[0]
    _strip_page_fields(title_first.header)
    _strip_page_fields(title_first.footer)
    _strip_page_fields(title_first.first_page_header)
    _strip_page_fields(title_first.first_page_footer)
    _strip_page_fields(title_first.even_page_header)
    _strip_page_fields(title_first.even_page_footer)
    merged_doc.save(output_path)


@api_view(['POST'])
@permission_classes([AllowAny])
def prepend_title_to_result(request, doc_id):
    """
    POST /api/prepend-title/<doc_id>/
    Берёт готовый титульник из файла title_file и возвращает объединённый docx:
    [титульник] + [исправленный отчёт].
    """
    template_path = None
    output_title = None
    output_merged = None
    try:
        doc = Document.objects.get(id=doc_id)
        if doc.status != 'completed':
            return Response({'error': 'Документ ещё не обработан'}, status=400)
        if not hasattr(doc, 'result') or not doc.result.output_file:
            return Response({'error': 'Исправленный файл не найден'}, status=404)

        title_file = request.FILES.get('title_file')
        if not title_file:
            return Response({'error': 'Прикрепите готовый титульный лист (.docx)'}, status=400)
        if not title_file.name.endswith('.docx'):
            return Response({'error': 'Титульный лист должен быть .docx'}, status=400)
        fd_title, output_title = tempfile.mkstemp(suffix='.docx', prefix='title_ready_')
        os.close(fd_title)
        with open(output_title, 'wb') as dst:
            for chunk in title_file.chunks():
                dst.write(chunk)

        fd_merge, output_merged = tempfile.mkstemp(suffix='.docx', prefix='with_title_')
        os.close(fd_merge)
        report_path = doc.result.output_file.path
        _merge_title_and_report(output_title, report_path, output_merged)

        with open(output_merged, 'rb') as f:
            content = f.read()
        response = HttpResponse(
            content,
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        )
        response['Content-Disposition'] = _attachment_content_disposition(
            f"GOST_with_title_{doc.filename}"
        )
        return response
    except Document.DoesNotExist:
        return Response({'error': 'Документ не найден'}, status=404)
    except Exception as e:
        return Response({'error': str(e)}, status=500)
    finally:
        for p in (template_path, output_title, output_merged):
            if p and os.path.exists(p):
                try:
                    os.unlink(p)
                except OSError:
                    pass

@csrf_exempt
def generate_title_page(request):
    """
    POST /api/title-page/
    Генерирует титульный лист и возвращает .docx файл.
    """
    if request.method != 'POST':
        return JsonResponse({'error': 'Только POST'}, status=405)

    template_path = None
    output_path = None

    try:
        data = {
            'institution':    request.POST.get('institution', ''),
            'work_title':     request.POST.get('work_title', 'ОТЧЕТ ПО ПРАКТИЧЕСКОЙ РАБОТЕ'),
            'work_number':    request.POST.get('work_number', ''),
            'specialty_code': request.POST.get('specialty_code', ''),
            'specialty_name': request.POST.get('specialty_name', ''),
            'subject':        request.POST.get('subject', ''),
            'group':          request.POST.get('group', ''),
            'student_id':     request.POST.get('student_id', ''),
            'student_name':   request.POST.get('student_name', ''),
            'teacher_name':   request.POST.get('teacher_name', ''),
            'city':           request.POST.get('city', 'Красноярск'),
            'year':           request.POST.get('year', '2026'),
        }

        # Рамка (необязательно)
        template_file = request.FILES.get('template')
        if template_file:
            if not template_file.name.endswith('.docx'):
                return JsonResponse({'error': 'Рамка должна быть .docx'}, status=400)
            import tempfile as _tf
            tmp = _tf.NamedTemporaryFile(delete=False, suffix='.docx')
            for chunk in template_file.chunks():
                tmp.write(chunk)
            tmp.close()
            template_path = tmp.name

        # Генерация
        import tempfile as _tf
        fd, output_path = _tf.mkstemp(suffix='.docx', prefix='title_')
        os.close(fd)

        from core.title_page import create_title_page
        create_title_page(data, output_path, template_path)

        # Читаем файл и возвращаем
        with open(output_path, 'rb') as f:
            content = f.read()

        from django.http import HttpResponse as HR
        response = HR(
            content,
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        work_num = data.get('work_number', '')
        fname = f'Титульный_лист_N{work_num}.docx' if work_num else 'Титульный_лист.docx'
        response['Content-Disposition'] = _attachment_content_disposition(fname)
        return response

    except Exception as e:
        import traceback
        traceback.print_exc()
        return JsonResponse({'error': str(e)}, status=500)

    finally:
        if template_path and os.path.exists(template_path):
            try:
                os.unlink(template_path)
            except:
                pass
        if output_path and os.path.exists(output_path):
            try:
                os.unlink(output_path)
            except:
                pass

def _update_error_stats(report, work_type):
    """Обновляет агрегированную статистику ошибок"""
    if not isinstance(report, dict):
        return

    errors = report.get('errors', [])
    for err in errors:
        err_type = err.get('type', 'unknown')
        stat, _ = ErrorStatistic.objects.get_or_create(
            error_type=err_type,
            work_type=work_type,
            defaults={'description': err.get('description', '')}
        )
        stat.count += 1
        stat.save()